"""
原版本的导出函数 - 直接复制
"""

from io import BytesIO
import pandas as pd
from datetime import datetime
import re
import logging
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

def generate_excel_in_memory(papers_data):
    """在内存中生成Excel文件"""
    try:
        papers = papers_data.get('papers', [])
        query = papers_data.get('query', '未知查询')

        # 创建DataFrame
        df_data = []
        for i, paper in enumerate(papers, 1):
            # 兼容两种数据格式：嵌套journal_info和平级字段
            journal_info = paper.get('journal_info', {})
            
            # 处理JCR分区
            jcr_quartile = journal_info.get('jcr_quartile', '') or paper.get('jcr_quartile', '')
            jcr_display = f"Q{jcr_quartile}" if jcr_quartile else ''

            # 处理中科院分区
            cas_quartile = journal_info.get('cas_quartile', '') or paper.get('cas_quartile', '')
            cas_display = f"{cas_quartile}区" if cas_quartile else ''

            # 获取相关度分数，支持多种字段名
            relevance_score = paper.get('relevance_score', 0) or paper.get('relevance', 0)
            if isinstance(relevance_score, (int, float)):
                relevance_display = f"{relevance_score * 100:.1f}%" if relevance_score <= 1 else f"{relevance_score:.1f}%"
            else:
                relevance_display = "0.0%"

            df_data.append({
                '序号': i,
                '标题': paper.get('title', ''),
                '作者': ', '.join(paper.get('authors', [])),
                '期刊': journal_info.get('title', '') or paper.get('journal', ''),
                '发表年份': paper.get('pub_year', ''),
                '影响因子': journal_info.get('impact_factor', '') or paper.get('impact_factor', ''),
                'JCR分区': jcr_display,
                '中科院分区': cas_display,
                '相关度': relevance_display,
                'PMID': paper.get('pmid', ''),
                'DOI': paper.get('doi', ''),
                '摘要': paper.get('abstract', '')
            })

        df = pd.DataFrame(df_data)

        # 使用BytesIO创建内存中的Excel文件
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='文献数据', index=False)

            # 获取工作表并设置格式
            worksheet = writer.sheets['文献数据']

            # 设置列宽
            column_widths = {
                'A': 8,   # 序号
                'B': 50,  # 标题
                'C': 30,  # 作者
                'D': 25,  # 期刊
                'E': 12,  # 年份
                'F': 12,  # 影响因子
                'G': 12,  # JCR分区
                'H': 12,  # 中科院分区
                'I': 12,  # 相关度
                'J': 15,  # PMID
                'K': 20,  # DOI
                'L': 80   # 摘要
            }

            for col, width in column_widths.items():
                worksheet.column_dimensions[col].width = width

        output.seek(0)
        file_content = output.getvalue()

        # 生成文件名 - 按原版本格式
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_query = re.sub(r'[^\w\u4e00-\u9fff]', '_', query)
        if len(safe_query) > 50:
            safe_query = safe_query[:50]
        filename = f'papers_{safe_query}__{timestamp}.xlsx'

        return file_content, filename

    except Exception as e:
        logger.error(f"生成Excel文件失败: {str(e)}")
        raise


def generate_word_in_memory(papers_data):
    """在内存中生成Word文件"""
    try:
        papers = papers_data.get('papers', [])
        query = papers_data.get('query', '未知查询')

        # 创建Word文档
        doc = Document()

        # 添加标题 - 按原版本格式
        title = doc.add_heading('文献检索报告', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加检索策略 - 按原版本格式
        doc.add_heading('检索策略', level=1)
        doc.add_paragraph(f'检索词：\n{query}')
        doc.add_paragraph(f'检索时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'检索到的文献数量：{len(papers)}')

        # 添加文献列表
        doc.add_heading('文献列表', level=1)

        for i, paper in enumerate(papers, 1):
            # 获取期刊信息 - 兼容两种数据格式
            journal_info = paper.get('journal_info', {})
            
            # 添加文献标题 - 按原版本格式
            p = doc.add_paragraph()
            p.add_run(f'{i}. ').bold = True
            title_run = p.add_run(paper.get('title', 'N/A'))
            title_run.bold = True
            
            # 添加作者信息
            authors_str = ', '.join(paper.get('authors', [])) if paper.get('authors') else 'N/A'
            doc.add_paragraph(f'作者：{authors_str}')
            
            # 添加期刊信息
            journal_title = journal_info.get('title', '') or paper.get('journal', '') or 'N/A'
            doc.add_paragraph(f'期刊：{journal_title}')
            doc.add_paragraph(f'发表时间：{paper.get("pub_year", "N/A")}')
            
            # 添加影响因子和分区信息
            impact_factor = journal_info.get('impact_factor', '') or paper.get('impact_factor', '') or 'N/A'
            doc.add_paragraph(f'影响因子：{impact_factor}')
            
            jcr_quartile = journal_info.get('jcr_quartile', '') or paper.get('jcr_quartile', '')
            jcr_display = f'Q{jcr_quartile}' if jcr_quartile else 'N/A'
            doc.add_paragraph(f'JCR分区：{jcr_display}')
            
            cas_quartile = journal_info.get('cas_quartile', '') or paper.get('cas_quartile', '')
            cas_display = f'{cas_quartile}区' if cas_quartile else 'N/A'
            doc.add_paragraph(f'CAS分区：{cas_display}')
            
            # 添加关键词信息
            keywords_str = ', '.join(paper.get('keywords', [])) if paper.get('keywords') else 'N/A'
            doc.add_paragraph(f'关键词：{keywords_str}')
            
            # 添加DOI和PMID
            doc.add_paragraph(f'DOI：{paper.get("doi", "N/A")}')
            doc.add_paragraph(f'PMID：{paper.get("pmid", "N/A")}')
            
            # 添加摘要 - 按原版本格式（斜体）
            if paper.get('abstract'):
                doc.add_paragraph('摘要：').add_run(paper['abstract']).italic = True
            
            # 添加分隔线 - 按原版本格式
            if i < len(papers):
                doc.add_paragraph('_' * 50)

        # 保存到内存
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        file_content = output.getvalue()

        # 生成文件名 - 按原版本格式
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_query = re.sub(r'[^\w\u4e00-\u9fff]', '_', query[:50])  # 限制查询长度为50个字符
        filename = f'papers_{safe_query}__{timestamp}.docx'

        return file_content, filename

    except Exception as e:
        logger.error(f"生成Word文件失败: {str(e)}")
        raise