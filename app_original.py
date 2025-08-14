from flask import Flask, request, jsonify, render_template, send_from_directory, make_response, send_file, session
import requests
import json
import os
from dotenv import load_dotenv
import nltk
from nltk.tokenize import sent_tokenize
from concurrent.futures import ThreadPoolExecutor, as_completed
import time
from bs4 import BeautifulSoup
import re
import logging
from datetime import datetime
import traceback
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
import base64
from langchain_community.retrievers.pubmed import PubMedRetriever
from langchain_community.document_loaders.pubmed import PubMedLoader
import sys
import urllib.parse
from typing import List, Dict, Optional, Union
from sklearn.feature_extraction.text import TfidfVectorizer
import numpy as np
from nltk.corpus import wordnet
import codecs
from journal_analyzer import JournalAnalyzer
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil
from pathlib import Path
import threading
import copy
from flask_socketio import SocketIO, emit
import asyncio
from flask_socketio import join_room, leave_room

# 确保必要的目录存在
for directory in ['logs', 'exports', 'static/images']:
    os.makedirs(directory, exist_ok=True)

# 确保NLTK数据包已下载
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
    nltk.data.find('corpora/wordnet')
except LookupError:
    print("正在下载NLTK数据包...")
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
    nltk.download('wordnet', quiet=True)
    print("NLTK数据包下载完成")

# 创建应用实例
app = Flask(__name__, template_folder='templates')
# 配置Socket.IO以支持更高并发
socketio = SocketIO(
    app,
    cors_allowed_origins="*",
    async_mode='threading',
    ping_timeout=60,
    ping_interval=25,
    max_http_buffer_size=1e8,  # 100MB
    manage_session=True,  # 启用Socket.IO的会话管理
    logger=True,  # 启用详细日志
    engineio_logger=True  # 启用Engine.IO日志
)

# 用户会话管理
active_users = {}  # 存储活跃用户的会话信息
user_stats_lock = threading.Lock()  # 用于保护用户统计数据的锁
user_messages = {}  # 存储每个用户的消息历史

# 文献数据缓存（内存存储）
papers_cache = {}  # 存储每个会话的文献数据，用于动态生成导出文件
papers_cache_lock = threading.Lock()  # 保护文献缓存的锁

# 嵌入模型相关度计算配置
EMBEDDING_BATCH_SIZE = 50  # 嵌入模型批量处理大小（优化为50篇提高效率）
EMBEDDING_MAX_TEXT_LENGTH = 1000  # 单个文本最大长度（避免API限制）

# 性能监控相关变量
performance_stats = {
    'response_times': [],
    'error_count': 0,
    'request_count': 0,
    'avg_response_time': 0.0
}
performance_stats_lock = threading.Lock()

# 用户统计相关变量
user_stats = {
    'total_visits': 0,
    'concurrent_users': 0,
    'peak_concurrent_users': 0,
    'visit_times': [],
    'hourly_stats': {str(i): 0 for i in range(24)}
}

# 线程池执行器
executor = ThreadPoolExecutor(
    max_workers=50,
    thread_name_prefix='NNScholar'
)

@socketio.on('connect')
def handle_connect():
    """处理用户连接"""
    try:
        # 从查询参数中获取会话ID
        session_id = request.args.get('sessionId')
        if not session_id:
            logger.warning("连接请求没有会话ID")
            return False
            
        # 将客户端加入对应的房间
        join_room(session_id)
        
        # 初始化用户会话
        with user_stats_lock:
            active_users[session_id] = {
                'connect_time': datetime.now(),
                'last_active': datetime.now(),
                'request_count': 0,
                'last_request_time': None
            }
            
        logger.info(f"用户连接成功 {session_id}")
        
        # 发送连接成功消息
        emit('search_progress', {
            'stage': 'connect',
            'message': '连接成功',
            'percentage': 0
        }, room=session_id)
        
        return True
        
    except Exception as e:
        logger.error(f"处理连接时出错: {str(e)}\n{traceback.format_exc()}")
        return False

@socketio.on('disconnect')
def handle_disconnect():
    """处理用户断开连接"""
    try:
        session_id = request.args.get('sessionId')
        if session_id:
            # 将客户端从房间中移除
            leave_room(session_id)
            
            with user_stats_lock:
                if session_id in active_users:
                    del active_users[session_id]
                    user_stats['concurrent_users'] = len(active_users)
            logger.info(f"用户断开 {session_id}, 当前在线人数: {user_stats['concurrent_users']}")
    except Exception as e:
        logger.error(f"处理断开连接时出错: {str(e)}\n{traceback.format_exc()}")

def update_user_stats(session_id, action='connect'):
    """更新用户统计信息"""
    try:
        with user_stats_lock:
            current_time = datetime.now()
            
            if action == 'connect':
                # 检查会话是否已存在
                if session_id not in active_users:
                    active_users[session_id] = {
                        'connect_time': current_time,
                        'last_active': current_time,
                        'request_count': 0,
                        'last_request_time': None,
                        'messages': []
                    }
                else:
                    # 更新现有会话
                    active_users[session_id]['last_active'] = current_time
                    
                # 更新统计信息
                user_stats['concurrent_users'] = len(active_users)
                if user_stats['concurrent_users'] > user_stats['peak_concurrent_users']:
                    user_stats['peak_concurrent_users'] = user_stats['concurrent_users']
                    
            elif action == 'disconnect':
                if session_id in active_users:
                    del active_users[session_id]
                    user_stats['concurrent_users'] = len(active_users)
                    
    except Exception as e:
        logger.error(f"更新用户统计时出错: {str(e)}")

def monitor_system_performance():
    """监控系统性能"""
    while True:
        try:
            with performance_stats_lock:
                if performance_stats['response_times']:
                    avg_response_time = sum(performance_stats['response_times']) / len(performance_stats['response_times'])
                    performance_stats['avg_response_time'] = avg_response_time
                    logger.info(f"平均响应时间: {avg_response_time:.2f}秒")
                    
                    if len(performance_stats['response_times']) > 1000:
                        performance_stats['response_times'] = performance_stats['response_times'][-1000:]
                
                if performance_stats['request_count'] > 0:
                    error_rate = (performance_stats['error_count'] / performance_stats['request_count']) * 100
                    logger.info(f"错误率: {error_rate:.2f}%")
                
                logger.info(f"当前在线用户: {user_stats['concurrent_users']}")
                logger.info(f"历史峰值: {user_stats['peak_concurrent_users']}")
                
                performance_stats['error_count'] = 0
                performance_stats['request_count'] = 0
        except Exception as e:
            logger.error(f"性能监控出错: {str(e)}")
        time.sleep(300)

@app.before_request
def before_request():
    """请求预处理"""
    try:
        request.start_time = time.time()
        session_id = request.headers.get('sid')
        if session_id:
            with user_stats_lock:
                if session_id in active_users:
                    active_users[session_id]['last_active'] = datetime.now()
    except Exception as e:
        logger.error(f"请求预处理时出错: {str(e)}")

@app.after_request
def after_request(response):
    """请求后处理"""
    try:
        with performance_stats_lock:
            if hasattr(request, 'start_time'):
                response_time = time.time() - request.start_time
                performance_stats['response_times'].append(response_time)
            
            performance_stats['request_count'] += 1
            if response.status_code >= 400:
                performance_stats['error_count'] += 1
    except Exception as e:
        logger.error(f"请求后处理时出错: {str(e)}")
    return response

def get_hourly_visits():
    """获取最近24小时的每小时访问量"""
    try:
        current_time = datetime.now()
        hourly_stats = [0] * 24
        
        for hour in range(24):
            hour_str = str(hour)
            if hour_str in user_stats['hourly_stats']:
                hour_index = (hour - current_time.hour) % 24
                hourly_stats[hour_index] = user_stats['hourly_stats'][hour_str]
        
        return hourly_stats
    except Exception as e:
        logger.error(f"获取小时访问量时出错: {str(e)}")
        return [0] * 24

def get_user_messages(session_id: str) -> List[Dict]:
    """
    获取用户的消息历史
    
    Args:
        session_id (str): 用户会话ID
        
    Returns:
        List[Dict]: 用户的消息历史列表
    """
    with user_stats_lock:
        return user_messages.get(session_id, [])

def handle_connect(session_id: str):
    """
    处理用户连接
    
    Args:
        session_id (str): 用户会话ID
    """
    with user_stats_lock:
        current_time = datetime.now()
        
        # 初始化或更新用户会话
        active_users[session_id] = {
            'connect_time': current_time,
            'last_active': current_time,
            'request_count': 0,
            'last_request_time': None,
            'messages': []
        }
        
        # 初始化用户消息列表
        if session_id not in user_messages:
            user_messages[session_id] = []
            
        # 更新用户统计信息
        user_stats['total_visits'] += 1
        user_stats['visit_times'].append(current_time)
        current_hour = str(current_time.hour)
        user_stats['hourly_stats'][current_hour] = user_stats['hourly_stats'].get(current_hour, 0) + 1
        user_stats['concurrent_users'] = len(active_users)
        
        # 更新峰值
        if user_stats['concurrent_users'] > user_stats['peak_concurrent_users']:
            user_stats['peak_concurrent_users'] = user_stats['concurrent_users']
            logger.info(f"新的并发峰值: {user_stats['peak_concurrent_users']}")

def handle_disconnect(session_id: str):
    """
    处理用户断开连接
    
    Args:
        session_id (str): 用户会话ID
    """
    with user_stats_lock:
        if session_id in active_users:
            del active_users[session_id]
        if session_id in user_messages:
            del user_messages[session_id]

def emit_to_user(session_id: str, event: str, data: Dict):
    """
    向指定用户发送消息
    
    Args:
        session_id (str): 用户会话ID
        event (str): 事件名称
        data (Dict): 要发送的数据
    """
    try:
        # 检查会话是否存在
        if session_id not in active_users:
            logger.warning(f"尝试向不存在的会话 {session_id} 发送消息")
            return
            
        # 发送消息
        try:
            logger.info(f"发送消息到用户 {session_id}: {event}")
            socketio.emit(event, data, room=session_id)
        except Exception as e:
            logger.error(f"发送消息到用户 {session_id} 失败: {str(e)}")
            raise
            
    except Exception as e:
        logger.error(f"处理消息发送时出错: {str(e)}")
        raise

def update_search_progress(session_id: str, stage: str, message: str, percentage: float = 0, **extra_data):
    """
    更新搜索进度
    
    Args:
        session_id (str): 用户会话ID
        stage (str): 当前阶段
        message (str): 进度消息
        percentage (float): 进度百分比
        **extra_data: 额外的数据字段
    """
    try:
        if not session_id:
            logger.error("无效的会话ID")
            return
            
        # 确保百分比在0-100之间
        percentage = min(100, max(0, percentage))
        
        # 构建进度数据
        progress_data = {
            'stage': stage,
            'message': message,
            'percentage': percentage,
            'timestamp': datetime.now().isoformat()
        }
        
        # 添加额外的数据字段
        progress_data.update(extra_data)
        
        # 记录日志
        logger.info(f"发送进度更新 [会话: {session_id}] - 阶段: {stage}, 消息: {message}, 进度: {percentage}%")
        
        try:
            # 发送进度更新
            socketio.emit('search_progress', progress_data, room=session_id)
        except Exception as e:
            logger.error(f"发送进度更新失败: {str(e)}")
            # 尝试发送错误消息
            try:
                socketio.emit('search_error', {'error': f"进度更新失败: {str(e)}"}, room=session_id)
            except:
                pass
                
    except Exception as e:
        logger.error(f"更新搜索进度时出错: {str(e)}\n{traceback.format_exc()}")
        try:
            socketio.emit('search_error', {'error': str(e)}, room=session_id)
        except:
            pass

def update_fetch_progress(session_id: str, stage: str, message: str, percentage: float, 
                         current: int = None, total: int = None, batch_info: Dict = None):
    """
    更新文献获取进度
    
    Args:
        session_id (str): 用户会话ID
        stage (str): 当前阶段
        message (str): 进度消息
        percentage (float): 进度百分比
        current (int, optional): 当前处理数量
        total (int, optional): 总数量
        batch_info (Dict, optional): 批次信息
    """
    try:
        progress_data = {
            'stage': stage,
            'message': message,
            'percentage': percentage
        }
        
        if current is not None:
            progress_data['current'] = current
            
        if total is not None:
            progress_data['total'] = total
            
        if batch_info:
            progress_data['batch_info'] = batch_info
            
        emit_to_user(session_id, 'fetch_progress', progress_data)
        
    except Exception as e:
        logger.error(f"更新文献获取进度时出错: {str(e)}")
        raise

# 创建必要的目录
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LOGS_DIR = os.path.join(BASE_DIR, 'logs')
EXPORTS_DIR = os.path.join(BASE_DIR, 'exports')
os.makedirs(LOGS_DIR, exist_ok=True)
os.makedirs(EXPORTS_DIR, exist_ok=True)

# 设置最大缓存时间（24小时，以秒为单位）
MAX_CACHE_TIME = 86400

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(os.path.join(LOGS_DIR, f'app_{datetime.now().strftime("%Y%m%d")}.log'), encoding='utf-8'),
        logging.StreamHandler(sys.stdout)  # 确保输出到标准输出
    ]
)
logger = logging.getLogger(__name__)

# 设置标准输出编码
sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer)
sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer)

# 检查.env文件
env_path = os.path.join(BASE_DIR, '.env')
if os.path.exists(env_path):
    logger.info(f"找到.env文件: {env_path}")
    with open(env_path, 'r', encoding='utf-8') as f:
        env_content = f.read()
        logger.info(f"环境文件内容预览 (前100字符): {env_content[:100]}...")
else:
    logger.error(f"未找到.env文件: {env_path}")

# 加载环境变量
load_dotenv(env_path, verbose=True, override=True)

def get_api_config() -> Dict[str, str]:
    """
    获取API配置并验证
    
    Returns:
        Dict[str, str]: 包含API配置的字典
    
    Raises:
        ValueError: 当缺少必要的环境变量时
    """
    # 直接从环境变量中读取并打印所有相关变量
    env_vars = {
        'DEEPSEEK_API_KEY': os.getenv('DEEPSEEK_API_KEY'),
        'DEEPSEEK_MODEL': os.getenv('DEEPSEEK_MODEL'),
        'DEEPSEEK_API_URL': os.getenv('DEEPSEEK_API_URL'),
        'PUBMED_API_KEY': os.getenv('PUBMED_API_KEY'),
        'PUBMED_EMAIL': os.getenv('PUBMED_EMAIL'),
        'TOOL_NAME': os.getenv('TOOL_NAME'),
        'PUBMED_API_URL': os.getenv('PUBMED_API_URL'),
    }
    
    logger.info("环境变量读取结果:")
    for key, value in env_vars.items():
        if 'API_KEY' in key and value:
            logger.info(f"{key}: {value[:4]}...{value[-4:]}")
        else:
            logger.info(f"{key}: {value}")
    
    config = {
        'deepseek_key': env_vars['DEEPSEEK_API_KEY'],
        'deepseek_model': env_vars['DEEPSEEK_MODEL'] or 'deepseek-chat',
        'deepseek_url': env_vars['DEEPSEEK_API_URL'] or 'https://api.deepseek.com/v1/chat/completions',
        'pubmed_key': env_vars['PUBMED_API_KEY'],
        'pubmed_email': env_vars['PUBMED_EMAIL'],
        'tool_name': env_vars['TOOL_NAME'] or 'nnscholar_pubmed',
        'pubmed_url': env_vars['PUBMED_API_URL'] or 'https://eutils.ncbi.nlm.nih.gov/entrez/eutils/'
    }
    
    # 验证必要的API密钥
    missing_keys = []
    if not config['deepseek_key']:
        missing_keys.append('DEEPSEEK_API_KEY')
    if not config['pubmed_key']:
        missing_keys.append('PUBMED_API_KEY')
    if not config['pubmed_email']:
        missing_keys.append('PUBMED_EMAIL')
        
    if missing_keys:
        raise ValueError(f"缺少必要的环境变量: {', '.join(missing_keys)}")
        
    return config

# 初始化API配置
try:
    API_CONFIG = get_api_config()
    DEEPSEEK_API_KEY = API_CONFIG['deepseek_key']
    DEEPSEEK_MODEL = API_CONFIG['deepseek_model']
    DEEPSEEK_API_URL = API_CONFIG['deepseek_url']
    PUBMED_API_KEY = API_CONFIG['pubmed_key']
    PUBMED_EMAIL = API_CONFIG['pubmed_email']
    TOOL_NAME = API_CONFIG['tool_name']
    PUBMED_BASE_URL = API_CONFIG['pubmed_url']

    # 嵌入模型配置
    EMBEDDING_API_KEY = os.getenv('EMBEDDING_API_KEY', '')
    EMBEDDING_API_URL = os.getenv('EMBEDDING_API_URL', 'https://api.siliconflow.cn/v1/embeddings')
    EMBEDDING_MODEL = os.getenv('EMBEDDING_MODEL', 'BAAI/bge-m3')
except Exception as e:
    logger.critical(f"API配置初始化失败: {str(e)}")
    raise

# 加载PubMed专家提示词模板
PROMPT_PATH = os.path.join(BASE_DIR, 'templates', 'pubmed_expert_prompt.md')
with open(PROMPT_PATH, 'r', encoding='utf-8') as f:
    EXPERT_PROMPT = f.read()

# 加载期刊数据
def load_journal_data():
    """加载期刊相关数据"""
    data_dir = os.path.join(BASE_DIR, 'data', 'journal_metrics')
    journal_data = {}
    if_trend_data = {}
    
    try:
        # 加载JCR和中科院分区数据
        jcr_file = os.path.join(data_dir, 'jcr_cas_ifqb.json')
        if not os.path.exists(jcr_file):
            logger.error(f"期刊数据文件不存在: {jcr_file}")
            return {}, {}
            
        logger.info(f"开始加载期刊数据文件: {jcr_file}")
        with open(jcr_file, 'r', encoding='utf-8') as f:
            try:
                journal_list = json.load(f)
                logger.info(f"成功加载期刊数据，包含 {len(journal_list)} 条记录")
                
                # 记录一些原始数据示例
                if len(journal_list) > 0:
                    sample_raw = journal_list[:3]
                    logger.info(f"原始数据示例: {json.dumps(sample_raw, ensure_ascii=False)}")
                
                for journal in journal_list:
                    # 处理ISSN和eISSN
                    issn = journal.get('issn', '').strip()
                    eissn = journal.get('eissn', '').strip()
                    
                    # 标准化ISSN格式（移除连字符）
                    issn = issn.replace('-', '') if issn else None
                    eissn = eissn.replace('-', '') if eissn else None
                    
                    # 使用所有可能的ISSN作为键
                    issns = [i for i in [issn, eissn] if i]
                    
                    if issns:
                        # 处理影响因子，确保是数值类型
                        impact_factor = journal.get('IF', 'N/A')
                        try:
                            if impact_factor != 'N/A':
                                impact_factor = float(impact_factor)
                        except (ValueError, TypeError):
                            impact_factor = 'N/A'
                            logger.warning(f"无效的影响因子值: {journal.get('IF')} for {journal.get('journal')}")
                        
                        journal_info = {
                            'title': journal.get('journal', ''),
                            'if': impact_factor,
                            'jcr_quartile': journal.get('Q', 'N/A'),
                            'cas_quartile': journal.get('B', 'N/A')
                        }
                        
                        # 为每个ISSN都存储期刊信息
                        for issn_key in issns:
                            journal_data[issn_key] = journal_info
                
                logger.info(f"成功加载 {len(journal_data)} 条期刊数据")
                # 记录一些转换后的数据示例
                if journal_data:
                    sample_converted = {k: journal_data[k] for k in list(journal_data.keys())[:3]}
                    logger.info(f"转换后的数据示例: {json.dumps(sample_converted, ensure_ascii=False)}")
            except json.JSONDecodeError as e:
                logger.error(f"期刊数据文件格式错误: {str(e)}")
                return {}, {}
        
        # 加载五年影响因子趋势数据
        trend_file = os.path.join(data_dir, '5year.json')
        if os.path.exists(trend_file):
            logger.info(f"开始加载影响因子趋势数据: {trend_file}")
            with open(trend_file, 'r', encoding='utf-8') as f:
                try:
                    if_trend_data = json.load(f)
                    if not isinstance(if_trend_data, dict):
                        logger.error("影响因子趋势数据格式错误：应为字典类型")
                        if_trend_data = {}
                    else:
                        logger.info(f"成功加载影响因子趋势数据，包含 {len(if_trend_data)} 条记录")
                except json.JSONDecodeError as e:
                    logger.error(f"影响因子趋势数据文件格式错误: {str(e)}")
                    if_trend_data = {}
        else:
            logger.warning(f"影响因子趋势数据文件不存在: {trend_file}")
        
        return journal_data, if_trend_data
        
    except Exception as e:
        logger.error(f"加载期刊数据失败: {str(e)}\n{traceback.format_exc()}")
        return {}, {}

# 全局变量
try:
    JOURNAL_DATA, IF_TREND_DATA = load_journal_data()
except Exception as e:
    logger.error(f"加载期刊数据失败: {str(e)}")
    JOURNAL_DATA, IF_TREND_DATA = {}, {}

def get_journal_metrics(issn):
    """获取期刊指标数据"""
    try:
        if not issn:
            logger.warning("ISSN为空")
            return None
            
        logger.info(f"开始获取期刊指标，ISSN: {issn}")
        
        if not isinstance(JOURNAL_DATA, dict):
            logger.error(f"期刊数据格式错误: {type(JOURNAL_DATA)}")
            return None
            
        if len(JOURNAL_DATA) == 0:
            logger.warning("期刊数据为空，请检查数据文件是否正确加载")
            return None
            
        # 标准化ISSN格式（移除连字符）
        issn = issn.replace('-', '')
        
        # 尝试直接获取
        journal_info = JOURNAL_DATA.get(issn)
        
        if not journal_info:
            # 尝试其他格式的ISSN
            issn_with_hyphen = f"{issn[:4]}-{issn[4:]}"
            journal_info = JOURNAL_DATA.get(issn_with_hyphen)
        
        if not journal_info:
            logger.warning(f"未找到ISSN对应的期刊信息: {issn}")
            return None
            
        logger.info(f"获取到的原始期刊信息: {json.dumps(journal_info, ensure_ascii=False)}")
        
        # 处理影响因子的显示格式
        impact_factor = journal_info.get('if', 'N/A')
        if isinstance(impact_factor, (int, float)):
            impact_factor = f"{impact_factor:.3f}"  # 格式化为三位小数
        
        metrics = {
            'title': journal_info.get('title', ''),
            'impact_factor': impact_factor,
            'jcr_quartile': journal_info.get('jcr_quartile', 'N/A'),
            'cas_quartile': journal_info.get('cas_quartile', 'N/A')
        }
        
        logger.info(f"处理后的期刊指标: {json.dumps(metrics, ensure_ascii=False)}")
        return metrics
        
    except Exception as e:
        logger.error(f"获取期刊指标时发生错误: {str(e)}\n{traceback.format_exc()}")
        return None

def get_if_trend(issn):
    """获取期刊近五年影响因子趋势"""
    if not issn or issn not in IF_TREND_DATA:
        return None
    
    trend_data = IF_TREND_DATA[issn]
    years = list(trend_data.keys())[-5:]
    ifs = [trend_data[year] for year in years]
    
    # 生成趋势图
    plt.figure(figsize=(8, 4))
    plt.plot(years, ifs, marker='o')
    plt.title('Impact Factor Trend (5 Years)')
    plt.xlabel('Year')
    plt.ylabel('Impact Factor')
    plt.grid(True)
    
    # 转换为base64图片
    buffer = BytesIO()
    plt.savefig(buffer, format='png')
    buffer.seek(0)
    image_png = buffer.getvalue()
    buffer.close()
    plt.close()
    
    return base64.b64encode(image_png).decode()

def filter_papers_by_metrics(papers, filters):
    """根据期刊指标筛选文献"""
    try:
        logger.info(f"开始筛选文献，筛选条件: {filters}")
        logger.info(f"待筛选文献数量: {len(papers)}")
        
        # 初始化统计信息
        stats = {
            'total': len(papers),
            'year_filtered': 0,
            'if_filtered': 0,
            'jcr_filtered': 0,
            'cas_filtered': 0,
            'final': 0
        }
        
        # 1. 年份筛选
        year_filtered = []
        if ('year_start' in filters and filters['year_start'] and 
            'year_end' in filters and filters['year_end']):
            year_start = int(filters['year_start'])
            year_end = int(filters['year_end'])
            logger.debug(f"应用年份筛选，范围: {year_start}-{year_end}")
            for paper in papers:
                pub_year = paper.get('pub_year')
                try:
                    pub_year = int(pub_year) if pub_year else None
                    if pub_year and year_start <= pub_year <= year_end:
                        year_filtered.append(paper)
                except (ValueError, TypeError) as e:
                    logger.warning(f"年份格式错误: {pub_year}, 错误信息: {str(e)}")
        else:
            year_filtered = papers.copy()
        stats['year_filtered'] = len(year_filtered)
        logger.info(f"1. 年份筛选 ({filters.get('year_start', '无')} - {filters.get('year_end', '无')}): {len(papers)} -> {len(year_filtered)}")
        
        # 2. 影响因子筛选
        if_filtered = []
        min_if = filters.get('impact_factor_min')
        max_if = filters.get('impact_factor_max')

        if min_if is not None or max_if is not None:
            logger.info(f"应用影响因子筛选，范围: {min_if} - {max_if}")
            for paper in year_filtered:
                journal_info = paper.get('journal_info', {})
                impact_factor = journal_info.get('impact_factor', 'N/A')
                try:
                    if impact_factor != 'N/A':
                        if isinstance(impact_factor, str):
                            impact_factor = float(impact_factor.replace(',', ''))
                        # 检查最小值
                        if min_if is not None and float(impact_factor) < min_if:
                            continue
                        # 检查最大值
                        if max_if is not None and float(impact_factor) > max_if:
                            continue
                        if_filtered.append(paper)
                    else:
                        # 没有影响因子信息的文献，如果没有设置筛选条件则保留
                        if min_if is None and max_if is None:
                            if_filtered.append(paper)
                except (ValueError, TypeError) as e:
                    logger.warning(f"影响因子格式错误: {impact_factor}, 错误信息: {str(e)}")
                    # 格式错误的文献，如果没有设置筛选条件则保留
                    if min_if is None and max_if is None:
                        if_filtered.append(paper)
        else:
            if_filtered = year_filtered.copy()
        stats['if_filtered'] = len(if_filtered)
        logger.info(f"2. 影响因子筛选 (>= {min_if or '无限制'}): {len(year_filtered)} -> {len(if_filtered)}")

        # 3. JCR分区筛选
        jcr_filtered = []
        jcr_quartiles = filters.get('jcr_quartiles', [])
        if jcr_quartiles:
            logger.info(f"应用JCR分区筛选，允许的分区: {jcr_quartiles}")
            for paper in if_filtered:
                journal_info = paper.get('journal_info', {})
                jcr_q = journal_info.get('jcr_quartile', 'N/A')
                if jcr_q != 'N/A' and jcr_q in jcr_quartiles:
                    jcr_filtered.append(paper)
                elif jcr_q == 'N/A':
                    # 没有JCR分区信息的文献也保留
                    jcr_filtered.append(paper)
        else:
            jcr_filtered = if_filtered.copy()
        stats['jcr_filtered'] = len(jcr_filtered)
        logger.info(f"3. JCR分区筛选 ({jcr_quartiles or '无限制'}): {len(if_filtered)} -> {len(jcr_filtered)}")

        # 4. CAS分区筛选
        cas_filtered = []
        cas_zones = filters.get('cas_zones', [])
        if cas_zones:
            logger.info(f"应用中科院分区筛选，允许的分区: {cas_zones}")
            cas_filters = [str(q) for q in cas_zones]
            for paper in jcr_filtered:
                journal_info = paper.get('journal_info', {})
                cas_q = journal_info.get('cas_quartile', 'N/A')
                if cas_q != 'N/A':
                    # 处理可能的格式（如B1 -> 1）
                    if isinstance(cas_q, str) and cas_q.startswith('B'):
                        cas_q = cas_q[1:]
                    if str(cas_q) in cas_filters:
                        cas_filtered.append(paper)
                else:
                    # 没有CAS分区信息的文献也保留
                    cas_filtered.append(paper)
        else:
            cas_filtered = jcr_filtered.copy()
        stats['cas_filtered'] = len(cas_filtered)
        logger.info(f"4. CAS分区筛选 ({cas_zones or '无限制'}): {len(jcr_filtered)} -> {len(cas_filtered)}")

        # 5. 文献类型筛选
        type_filtered = []
        publication_types = filters.get('publication_types', [])
        if publication_types:
            logger.info(f"应用文献类型筛选，允许的类型: {publication_types}")
            for paper in cas_filtered:
                paper_types = paper.get('publication_types', [])
                # 检查是否有任何匹配的文献类型
                if any(ptype in publication_types for ptype in paper_types):
                    type_filtered.append(paper)
                elif not paper_types:
                    # 没有类型信息的文献也保留
                    type_filtered.append(paper)
        else:
            type_filtered = cas_filtered.copy()
        stats['type_filtered'] = len(type_filtered)
        logger.info(f"5. 文献类型筛选 ({publication_types or '无限制'}): {len(cas_filtered)} -> {len(type_filtered)}")

        # 6. 语言筛选
        lang_filtered = []
        languages = filters.get('languages', [])
        if languages and 'other' not in languages:  # 如果选择了"其他语言"则不筛选
            logger.info(f"应用语言筛选，允许的语言: {languages}")
            for paper in type_filtered:
                paper_lang = paper.get('language', 'eng')  # 默认为英文
                if paper_lang in languages:
                    lang_filtered.append(paper)
                elif not paper_lang:
                    # 没有语言信息的文献也保留
                    lang_filtered.append(paper)
        else:
            lang_filtered = type_filtered.copy()
        stats['lang_filtered'] = len(lang_filtered)
        logger.info(f"6. 语言筛选 ({languages or '无限制'}): {len(type_filtered)} -> {len(lang_filtered)}")

        # 7. 计算综合得分并排序
        # 单句模式下，相关性权重为0.7，影响因子权重为0.3
        for paper in lang_filtered:
            relevance = float(paper.get('relevance', 0))
            journal_info = paper.get('journal_info', {})
            impact_factor = journal_info.get('impact_factor', 'N/A')
            
            try:
                if impact_factor != 'N/A':
                    if isinstance(impact_factor, str):
                        impact_factor = float(impact_factor.replace(',', ''))
                    # 将影响因子归一化到0-100的范围（假设最高影响因子为50）
                    if_score = min(100, (float(impact_factor) / 50) * 100)
                else:
                    if_score = 0
            except (ValueError, TypeError):
                if_score = 0
            
            # 计算综合得分
            paper['composite_score'] = (relevance * 0.7) + (if_score * 0.3)
            logger.debug(f"文献 {paper.get('pmid')} 的综合得分: {paper['composite_score']:.1f} (相关性: {relevance:.1f}, IF得分: {if_score:.1f})")
        
        # 按综合得分排序
        filtered_papers = sorted(
            lang_filtered,
            key=lambda x: x.get('composite_score', 0),
            reverse=True
        )
        
        # 限制返回数量
        try:
            papers_limit = int(filters.get('papers_limit', 200))  # 确保转换为整数
        except (ValueError, TypeError) as e:
            logger.warning(f"无效的papers_limit值，使用默认值10。错误: {str(e)}")
            papers_limit = 10
        papers = filtered_papers[:papers_limit]
        
        stats['final'] = len(papers)
        
        # 输出详细的筛选统计信息
        logger.info("\n筛选过程统计:")
        logger.info(f"初始文献数量: {stats['total']}")
        logger.info(f"1. 年份筛选后: {stats['year_filtered']} 篇")
        logger.info(f"2. 影响因子筛选后: {stats['if_filtered']} 篇")
        logger.info(f"3. JCR分区筛选后: {stats['jcr_filtered']} 篇")
        logger.info(f"4. CAS分区筛选后: {stats['cas_filtered']} 篇")
        logger.info(f"5. 文献类型筛选后: {stats['type_filtered']} 篇")
        logger.info(f"6. 语言筛选后: {stats['lang_filtered']} 篇")
        logger.info(f"7. 最终结果: {stats['final']} 篇")
        
        return papers, stats
        
    except Exception as e:
        logger.error(f"筛选文献时发生错误: {str(e)}\n{traceback.format_exc()}")
        raise

def handle_api_error(func):
    """API错误处理装饰器"""
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except requests.exceptions.RequestException as e:
            logger.error(f"API请求错误: {str(e)}")
            return jsonify({
                'status': 'error',
                'message': '外部API请求失败，请稍后重试'
            }), 503
        except Exception as e:
            logger.error(f"未预期的错误: {str(e)}\n{traceback.format_exc()}")
            return jsonify({
                'status': 'error',
                'message': '服务器内部错误'
            }), 500
    wrapper.__name__ = func.__name__  # 保留原函数名
    return wrapper

def call_deepseek_api(prompt):
    """调用DeepSeek API进行文本处理"""
    # 记录API密钥前三位用于调试
    if DEEPSEEK_API_KEY:
        logger.info(f"DeepSeek API密钥前三位: {DEEPSEEK_API_KEY[:3]}...")
    else:
        logger.error("DeepSeek API密钥未设置")
        
    headers = {
        'Authorization': f'Bearer {DEEPSEEK_API_KEY}',
        'Content-Type': 'application/json'
    }
    
    data = {
        'model': DEEPSEEK_MODEL,
        'messages': [
            {'role': 'system', 'content': EXPERT_PROMPT},
            {'role': 'user', 'content': prompt}
        ]
    }
    
    try:
        logger.info(f"调用DeepSeek API，模型: {DEEPSEEK_MODEL}, URL: {DEEPSEEK_API_URL}")
        response = requests.post(
            DEEPSEEK_API_URL,
            headers=headers,
            json=data
        )
        
        # 检查HTTP状态码
        response.raise_for_status()
        
        # 解析JSON响应
        response_data = response.json()
        
        # 记录完整响应用于调试
        logger.debug(f"DeepSeek API响应: {response_data}")
        
        # 验证响应格式
        if 'choices' not in response_data:
            error_msg = response_data.get('error', {}).get('message', '未知错误')
            logger.error(f"DeepSeek API返回格式错误: {error_msg}")
            raise ValueError(f"DeepSeek API错误: {error_msg}")
            
        return response_data['choices'][0]['message']['content']
        
    except requests.exceptions.RequestException as e:
        logger.error(f"调用DeepSeek API时发生网络错误: {str(e)}")
        raise
    except ValueError as e:
        logger.error(f"处理DeepSeek API响应时发生错误: {str(e)}")
        raise
    except Exception as e:
        logger.error(f"调用DeepSeek API时发生未预期的错误: {str(e)}")
        raise

def call_embedding_api(texts):
    """调用配置的嵌入模型API获取文本向量"""
    try:
        # 检查API配置
        if not EMBEDDING_API_KEY or not EMBEDDING_API_URL:
            logger.warning("嵌入模型API配置不完整，使用备用方法")
            return call_fallback_embedding(texts)

        # 确保texts是列表格式
        if isinstance(texts, str):
            texts = [texts]

        headers = {
            'Authorization': f'Bearer {EMBEDDING_API_KEY}',
            'Content-Type': 'application/json'
        }

        data = {
            'model': EMBEDDING_MODEL,
            'input': texts
        }

        logger.info(f"调用嵌入模型API ({EMBEDDING_MODEL})，处理 {len(texts)} 个文本")

        response = requests.post(
            EMBEDDING_API_URL,
            headers=headers,
            json=data,
            timeout=30
        )

        if response.status_code == 200:
            result = response.json()
            embeddings = [item['embedding'] for item in result['data']]
            logger.info(f"成功获取 {len(embeddings)} 个嵌入向量")
            return embeddings
        else:
            logger.error(f"嵌入模型API调用失败: {response.status_code} - {response.text}")
            logger.info("回退到备用嵌入方法")
            return call_fallback_embedding(texts)

    except Exception as e:
        logger.error(f"调用嵌入模型API时出错: {str(e)}")
        logger.info("回退到备用嵌入方法")
        return call_fallback_embedding(texts)

def call_fallback_embedding(texts):
    """备用嵌入方法：使用文本特征生成向量"""
    try:
        logger.info(f"使用备用文本特征方法，处理 {len(texts)} 个文本")

        embeddings = []
        for text in texts:
            embedding = create_text_features(text)
            embeddings.append(embedding)

        logger.info(f"成功生成 {len(embeddings)} 个特征向量")
        return embeddings

    except Exception as e:
        logger.error(f"备用嵌入方法失败: {str(e)}")
        return None

def create_text_features(text):
    """创建基于文本特征的向量表示"""
    try:
        import hashlib
        import numpy as np

        # 文本预处理
        text = text.lower().strip()
        words = text.split()

        # 创建基于文本特征的向量（384维，模拟常见嵌入模型维度）
        vector_size = 384
        vector = np.zeros(vector_size)

        # 基于词汇特征
        for i, word in enumerate(words[:50]):  # 最多处理前50个词
            # 使用词的哈希值作为特征
            word_hash = int(hashlib.md5(word.encode()).hexdigest()[:8], 16)
            idx = word_hash % vector_size
            vector[idx] += 1.0 / (i + 1)  # 位置权重

        # 基于字符特征
        for i, char in enumerate(text[:200]):  # 最多处理前200个字符
            if char.isalnum():
                char_hash = ord(char)
                idx = char_hash % vector_size
                vector[idx] += 0.1

        # 基于文本长度特征
        length_features = [
            len(text) / 1000.0,  # 文本长度
            len(words) / 100.0,  # 词数
            len(set(words)) / len(words) if words else 0,  # 词汇多样性
        ]

        # 将长度特征添加到向量的前几个位置
        for i, feature in enumerate(length_features[:3]):
            if i < vector_size:
                vector[i] += feature

        # 归一化向量
        norm = np.linalg.norm(vector)
        if norm > 0:
            vector = vector / norm

        return vector.tolist()

    except Exception as e:
        logger.error(f"创建文本特征时出错: {str(e)}")
        # 返回随机向量作为备用
        import random
        return [random.random() - 0.5 for _ in range(384)]

def calculate_cosine_similarity(vec1, vec2):
    """计算两个向量的余弦相似度"""
    try:
        import numpy as np

        # 转换为numpy数组
        vec1 = np.array(vec1)
        vec2 = np.array(vec2)

        # 计算余弦相似度
        dot_product = np.dot(vec1, vec2)
        norm1 = np.linalg.norm(vec1)
        norm2 = np.linalg.norm(vec2)

        if norm1 == 0 or norm2 == 0:
            return 0.0

        similarity = dot_product / (norm1 * norm2)
        return float(similarity)

    except Exception as e:
        logger.error(f"计算余弦相似度时出错: {str(e)}")
        return 0.0

def parse_pubmed_xml(xml_content):
    """解析PubMed XML响应"""
    logger.info("开始解析PubMed XML响应")
    soup = BeautifulSoup(xml_content, 'lxml')
    articles = []
    article_count = len(soup.find_all('pubmedarticle'))
    logger.info(f"找到 {article_count} 篇文章记录")
    
    for i, article in enumerate(soup.find_all('pubmedarticle'), 1):
        try:
            logger.info(f"开始解析第 {i}/{article_count} 篇文章")
            
            # 提取文章标题
            title = article.find('articletitle')
            title = title.text if title else 'No title available'
            # logger.info(f"文章标题: {title[:100]}...")
            
            # 提取发表年份
            pub_date = article.find('pubdate')
            pub_year = None
            if pub_date:
                # 尝试从Year标签提取
                year_elem = pub_date.find('year')
                if year_elem and year_elem.text:
                    try:
                        pub_year = int(year_elem.text)
                        # logger.info(f"成功提取发表年份: {pub_year}")
                    except ValueError:
                        logger.warning(f"无效的年份格式: {year_elem.text}")
                else:
                    # 尝试从MedlineDate中提取
                    medline_date = pub_date.find('medlinedate')
                    if medline_date and medline_date.text:
                        try:
                            # 提取第一个四位数字作为年份
                            year_match = re.search(r'\b\d{4}\b', medline_date.text)
                            if year_match:
                                pub_year = int(year_match.group())
                                # logger.info(f"从MedlineDate提取到年份: {pub_year}")
                        except ValueError:
                            logger.warning(f"无法从MedlineDate提取年份: {medline_date.text}")
            
            # 提取期刊信息
            journal = article.find('journal')
            journal_info = {}
            if journal:
                # 提取ISSN
                issn_elem = journal.find('issn')
                if issn_elem:
                    issn = issn_elem.text
                    # logger.info(f"找到ISSN: {issn}")
                else:
                    issn = None
                    logger.warning("未找到ISSN")
                journal_info['issn'] = issn
                
                # 提取期刊标题
                journal_title = journal.find('title')
                journal_info['title'] = journal_title.text if journal_title else ''
                # logger.info(f"期刊标题: {journal_info['title']}")
                
                # 获取期刊指标
                if issn:
                    # logger.info(f"开始获取期刊 {issn} 的指标信息")
                    metrics = get_journal_metrics(issn)
                    if metrics:
                        # logger.info(f"成功获取期刊指标: {metrics}")
                        journal_info.update(metrics)
                    else:
                        logger.warning(f"未能获取期刊 {issn} 的指标信息")
                
            # 构建文章数据
            article_data = {
                'title': title,
                'abstract': article.find('abstract').text if article.find('abstract') else 'No abstract available',
                'authors': [f"{author.find('lastname').text} {author.find('forename').text}" 
                          for author in article.find('authorlist').find_all('author') 
                          if author.find('lastname') and author.find('forename')] if article.find('authorlist') else [],
                'pub_date': (lambda d: f"{d.find('year').text if d.find('year') else ''} {d.find('month').text if d.find('month') else ''}".strip())(article.find('pubdate')) if article.find('pubdate') else 'Date not available',
                'pub_year': pub_year,  # 确保年份被正确存储
                'pmid': article.find('pmid').text if article.find('pmid') else '',
                'url': f'https://pubmed.ncbi.nlm.nih.gov/{article.find("pmid").text}/' if article.find('pmid') else '#',
                'journal_info': journal_info,
                'journal_issn': journal_info.get('issn', '')
            }
            
            # 提取关键词
            keywords = []
            
            # 记录XML结构信息，帮助调试
            article_id = article.find('pmid').text if article.find('pmid') else 'unknown'
            logger.info(f"分析文章 {article_id} 的关键词结构")
            
            # 首先尝试提取<keywordlist owner="NOTNLM">标签下的关键词
            medline_citation = article.find('medlinecitation')
            if medline_citation:
                keyword_list = medline_citation.find('keywordlist', {'owner': 'NOTNLM'})
                if keyword_list:
                    keywords = [k.text.strip() for k in keyword_list.find_all('keyword')]
                    logger.info(f"从NOTNLM关键词列表中提取了 {len(keywords)} 个关键词")
                    if keywords:
                        logger.info(f"关键词示例: {keywords[:3]}")
            
            # 如果没有找到关键词，尝试其他可能的位置
            if not keywords:
                # 尝试查找任何keywordlist标签
                keyword_list = article.find('keywordlist') or article.find('KeywordList')
                if keyword_list:
                    keyword_elements = keyword_list.find_all('keyword') or keyword_list.find_all('Keyword')
                    keywords = [k.text.strip() for k in keyword_elements]
                    logger.info(f"成功提取 {len(keywords)} 个关键词，标签名: {keyword_list.name}")
                    if keywords:
                        logger.info(f"关键词示例: {keywords[:3]}")
                else:
                    # 尝试从MeSH中提取
                    mesh_list = article.find('meshheadinglist') or article.find('MeshHeadingList')
                    if not mesh_list and medline_citation:
                        mesh_list = medline_citation.find('meshheadinglist') or medline_citation.find('MeshHeadingList')
                    
                    if mesh_list:
                        mesh_terms = []
                        mesh_elements = mesh_list.find_all('meshheading') or mesh_list.find_all('MeshHeading')
                        for mesh in mesh_elements:
                            descriptor = mesh.find('descriptorname') or mesh.find('DescriptorName')
                            if descriptor:
                                mesh_terms.append(descriptor.text.strip())
                        if mesh_terms:
                            keywords = mesh_terms
                            logger.info(f"从MeSH中提取了 {len(keywords)} 个关键词")
                            if keywords:
                                logger.info(f"MeSH关键词示例: {keywords[:3]}")
                        else:
                            logger.info("MeSH列表为空")
                    else:
                        logger.info("未找到关键词列表或MeSH列表")
            
            article_data['keywords'] = keywords
            
            # 提取DOI
            article_id_list = article.find('articleidlist')
            if article_id_list:
                for article_id in article_id_list.find_all('articleid'):
                    if article_id.get('idtype') == 'doi':
                        article_data['doi'] = article_id.text
                        logger.info(f"成功提取DOI: {article_id.text}")
                        break
                if 'doi' not in article_data:
                    article_data['doi'] = ''
                    logger.warning(f"文献 {article_data['pmid']} 未找到DOI信息")
            else:
                article_data['doi'] = ''
                logger.warning(f"文献 {article_data['pmid']} 缺少ArticleIdList")
            
            logger.info(f"文章数据构建完成: PMID={article_data['pmid']}, 年份={article_data['pub_year']}, 关键词数量={len(keywords)}")
            articles.append(article_data)
            
        except Exception as e:
            logger.error(f"解析第 {i} 篇文章时发生错误: {str(e)}\n{traceback.format_exc()}")
            continue
    
    logger.info(f"完成XML解析，成功解析 {len(articles)}/{article_count} 篇文章")
    return articles

# 初始化全局变量
model = None

def preprocess_text(text):
    """文本预处理函数
    
    Args:
        text (str): 输入文本
        
    Returns:
        str: 预处理后的文本
    """
    if not text:
        return ""
        
    # 转换为小写
    text = text.lower()
    
    # 移除标点符号
    text = re.sub(r'[^\w\s]', ' ', text)
    
    # 移除多余空格
    text = re.sub(r'\s+', ' ', text)
    
    # 移除数字
    text = re.sub(r'\d+', '', text)
    
    return text.strip()



def calculate_relevance_improved(sentence, paper):
    """完全基于嵌入模型的相关性计算方法"""
    try:
        # 使用嵌入模型计算语义相似度
        embedding_score = calculate_embedding_relevance(sentence, paper)

        if embedding_score is not None:
            # 如果嵌入模型计算成功，使用嵌入分数
            final_score = embedding_score
            logger.debug(f"相关度计算结果（嵌入模型）: {final_score:.1f}")
        else:
            # 如果嵌入模型失败，返回默认低分
            logger.warning("嵌入模型计算失败，返回默认分数")
            final_score = 10.0  # 给一个较低的默认分数

        # 确保分数在0-100之间
        final_score = max(0.0, min(100.0, final_score))

        return round(final_score, 1)

    except Exception as e:
        logger.error(f"计算相关性时出错: {str(e)}")
        return 10.0  # 返回默认低分

def calculate_embedding_relevance(query, paper):
    """使用嵌入模型计算语义相关度"""
    try:
        # 准备文本
        query_text = str(query).strip()
        title = paper.get('title', '').strip()
        abstract = paper.get('abstract', '').strip()

        if not query_text or not title:
            logger.warning("查询文本或文献标题为空，无法计算嵌入相关度")
            return None

        # 构建文献文本（标题 + 摘要）
        paper_text = title
        if abstract:
            paper_text += " " + abstract

        # 调用嵌入模型API
        texts = [query_text, paper_text]
        embeddings = call_embedding_api(texts)

        if embeddings is None or len(embeddings) != 2:
            logger.warning("嵌入模型API调用失败或返回结果不完整")
            return None

        # 计算余弦相似度
        similarity = calculate_cosine_similarity(embeddings[0], embeddings[1])

        # 将相似度转换为0-100的分数
        # 余弦相似度范围是[-1, 1]，我们将其映射到[0, 100]
        score = (similarity + 1) * 50

        logger.info(f"嵌入模型相关度计算: 相似度={similarity:.3f}, 分数={score:.1f}")

        return score

    except Exception as e:
        logger.error(f"使用嵌入模型计算相关度时出错: {str(e)}")
        return None

def calculate_batch_embedding_relevance(query, papers, batch_size=40):
    """优化的批量计算多个文献的嵌入相关度"""
    try:
        query_text = str(query).strip()
        if not query_text:
            logger.warning("查询文本为空，无法计算批量嵌入相关度")
            return {}

        results = {}
        total_batches = (len(papers) + batch_size - 1) // batch_size

        logger.info(f"开始批量嵌入计算：{len(papers)} 篇文献，分 {total_batches} 批处理")

        # 分批处理文献
        for batch_idx in range(0, len(papers), batch_size):
            batch_papers = papers[batch_idx:batch_idx + batch_size]
            current_batch_num = batch_idx // batch_size + 1

            # 准备批量文本
            texts = [query_text]  # 第一个是查询文本
            paper_indices = []
            valid_papers = []

            for j, paper in enumerate(batch_papers):
                title = paper.get('title', '').strip()
                abstract = paper.get('abstract', '').strip()

                if title:  # 只处理有标题的文献
                    # 优化文本组合：限制长度避免API限制
                    paper_text = title
                    if abstract:
                        # 限制摘要长度，避免文本过长
                        abstract_truncated = abstract[:500] if len(abstract) > 500 else abstract
                        paper_text += " " + abstract_truncated

                    texts.append(paper_text)
                    paper_indices.append(batch_idx + j)  # 记录原始索引
                    valid_papers.append(paper)

            if len(texts) <= 1:  # 只有查询文本，没有有效的文献文本
                logger.warning(f"批次 {current_batch_num} 没有有效文献，跳过")
                continue

            # 调用嵌入模型API（添加重试机制）
            embeddings = None
            max_retries = 3

            for retry in range(max_retries):
                try:
                    embeddings = call_embedding_api(texts)
                    if embeddings is not None and len(embeddings) == len(texts):
                        break
                    else:
                        logger.warning(f"批次 {current_batch_num} 嵌入API返回不完整，重试 {retry + 1}/{max_retries}")
                except Exception as e:
                    logger.warning(f"批次 {current_batch_num} 嵌入API调用失败，重试 {retry + 1}/{max_retries}: {str(e)}")

                if retry < max_retries - 1:
                    time.sleep(1)  # 等待1秒后重试

            if embeddings is None or len(embeddings) != len(texts):
                logger.error(f"批次 {current_batch_num} 嵌入API调用最终失败，跳过该批次")
                continue

            # 计算相似度
            query_embedding = embeddings[0]

            for k, paper_idx in enumerate(paper_indices):
                try:
                    paper_embedding = embeddings[k + 1]  # k+1 因为第0个是查询嵌入

                    similarity = calculate_cosine_similarity(query_embedding, paper_embedding)

                    # 优化分数映射：使用更合理的映射函数
                    # 余弦相似度 [-1, 1] -> [0, 100]，但给予更好的分布
                    if similarity >= 0.8:
                        score = 85 + (similarity - 0.8) * 75  # 高相似度：85-100分
                    elif similarity >= 0.6:
                        score = 70 + (similarity - 0.6) * 75  # 中高相似度：70-85分
                    elif similarity >= 0.4:
                        score = 50 + (similarity - 0.4) * 100  # 中等相似度：50-70分
                    elif similarity >= 0.2:
                        score = 30 + (similarity - 0.2) * 100  # 中低相似度：30-50分
                    else:
                        score = max(10, (similarity + 1) * 20)  # 低相似度：10-30分

                    results[paper_idx] = round(min(100.0, max(0.0, score)), 1)

                except Exception as e:
                    logger.error(f"计算文献 {paper_idx} 相似度时出错: {str(e)}")
                    results[paper_idx] = 25.0  # 给一个默认分数

            logger.info(f"批次 {current_batch_num}/{total_batches} 完成，处理了 {len(paper_indices)} 篇文献")

        logger.info(f"批量嵌入相关度计算完成，共处理 {len(results)} 篇文献")
        return results

    except Exception as e:
        logger.error(f"批量计算嵌入相关度时出错: {str(e)}")
        return {}

def search_pubmed(query, max_results=600):
    """直接使用PubMed API搜索文献"""
    try:
        logger.info(f"开始PubMed搜索，检索策略: {query}, 最大结果数: {max_results}")
        
        session_id = request.headers.get('sid')
        if not session_id:
            raise ValueError("无效的会话ID")
        
        # 发送搜索开始信息
        update_search_progress(session_id, 'search_start', "开始PubMed搜索...", 0)
        
        # 使用提供的检索策略
        search_strategy = query

        # 构建PubMed搜索请求
        search_params = {
            'db': 'pubmed',
            'term': search_strategy,
            'retmax': str(max_results),
            'retmode': 'json',
            'api_key': PUBMED_API_KEY
        }
        
        # 发送正在搜索的信息
        update_search_progress(session_id, 'searching', "正在PubMed中搜索文献...", 60)
        
        # 发送搜索请求（添加重试机制）
        search_url = f"{PUBMED_BASE_URL}esearch.fcgi"

        max_retries = 5
        response = None

        # 创建一个更稳定的session
        session = requests.Session()
        session.headers.update({
            'User-Agent': 'NNScholar/1.0 (mailto:support@nnscholar.com)',
            'Accept': 'application/json',
            'Connection': 'keep-alive'
        })

        for attempt in range(max_retries):
            try:
                logger.info(f"PubMed搜索请求 (尝试 {attempt + 1}/{max_retries}): {search_url}")
                logger.info(f"搜索参数: {search_params}")

                response = session.get(
                    search_url,
                    params=search_params,
                    timeout=(10, 30),  # (连接超时, 读取超时)
                    verify=True,  # 验证SSL证书
                    allow_redirects=True
                )

                # 检查响应状态
                if response.status_code == 200:
                    logger.info(f"PubMed搜索请求成功 (尝试 {attempt + 1})")
                    break
                else:
                    logger.warning(f"PubMed返回HTTP {response.status_code} (尝试 {attempt + 1})")
                    if attempt == max_retries - 1:
                        update_search_progress(session_id, 'error', f"PubMed服务器错误: HTTP {response.status_code}", 100)
                        return [], search_strategy, 0, 0

            except (requests.exceptions.ConnectionError,
                    requests.exceptions.Timeout,
                    requests.exceptions.SSLError,
                    requests.exceptions.ChunkedEncodingError) as e:
                logger.warning(f"PubMed搜索请求失败 (尝试 {attempt + 1}/{max_retries}): {type(e).__name__}: {str(e)}")
                if attempt == max_retries - 1:
                    update_search_progress(session_id, 'error', f"网络连接失败: {str(e)}", 100)
                    return [], search_strategy, 0, 0

                # 指数退避，但不要等太久
                wait_time = min(2 ** attempt, 10)
                logger.info(f"等待 {wait_time} 秒后重试...")
                time.sleep(wait_time)

            except Exception as e:
                logger.error(f"PubMed搜索出现未预期错误 (尝试 {attempt + 1}): {str(e)}")
                if attempt == max_retries - 1:
                    update_search_progress(session_id, 'error', f"搜索出错: {str(e)}", 100)
                    return [], search_strategy, 0, 0
                time.sleep(1)
        
        if response.status_code != 200:
            logger.error(f"PubMed搜索请求失败: HTTP {response.status_code}")
            
            # 发送搜索失败信息
            update_search_progress(session_id, 'search_failed', f"PubMed搜索请求失败: HTTP {response.status_code}", 100)
            
            return [], search_strategy, 0, 0
            
        search_result = response.json()
        total_count = int(search_result.get('esearchresult', {}).get('count', 0))
        id_list = search_result.get('esearchresult', {}).get('idlist', [])
        
        # 发送搜索结果信息
        update_search_progress(session_id, 'search_complete', f"找到 {total_count} 篇文献", total_count, result_count=len(id_list))
        
        return id_list, search_strategy, total_count, len(id_list)
    except Exception as e:
        logger.error(f"PubMed搜索出错: {str(e)}\n{traceback.format_exc()}")
        
        # 发送错误信息
        if session_id:
            update_search_progress(session_id, 'error', f"搜索出错: {str(e)}", 100)
            
        return [], "", 0, 0

def extract_basic_terms(text):
    """提取基本关键词"""
    # 提取括号中的缩写
    abbreviations = re.findall(r'\(([A-Z]+)\)', text)
    
    # 提取主要医学术语
    key_terms = []
    medical_terms = [
        'Coronary computed tomography angiography',
        'CCTA',
        'atherosclerotic plaque',
        'coronary'
    ]
    
    for term in medical_terms:
        if term.lower() in text.lower():
            key_terms.append(term)
    
    # 合并缩写和关键词，限制数量
    all_terms = key_terms + abbreviations
    return list(set(all_terms))[:3]  # 最多返回3个关键词

def fetch_paper_details(id_list):
    """分批获取文献详细信息"""
    try:
        if not id_list:
            return []
            
        session_id = request.headers.get('sid')
        if not session_id:
            raise ValueError("无效的会话ID")
            
        # 将ID列表分成较小的批次，每批300个ID
        batch_size = 300
        all_papers = []
        total_batches = (len(id_list) + batch_size - 1) // batch_size
        
        logger.info(f"开始获取文献详情，共 {len(id_list)} 篇文献，分 {total_batches} 批处理")
        
        # 发送初始进度信息
        update_fetch_progress(session_id, 'fetch_start', f"开始获取文献详情，共 {len(id_list)} 篇文献，将分 {total_batches} 批处理", 0, total=len(id_list))
        
        for i in range(0, len(id_list), batch_size):
            batch_ids = id_list[i:i+batch_size]
            current_batch = i // batch_size + 1
            logger.info(f"正在处理第 {current_batch}/{total_batches} 批文献 ({len(batch_ids)} 篇)")
            
            # 发送批次开始处理的消息
            update_fetch_progress(session_id, 'batch_start', f"正在处理第 {current_batch}/{total_batches} 批文献 ({len(batch_ids)} 篇)", 
                                (current_batch - 1) / total_batches * 100,
                                total=total_batches, 
                                batch_info={'current_batch': current_batch, 'total_batches': total_batches, 'batch_size': len(batch_ids)})
            
            # 构建请求参数
            fetch_params = {
                'db': 'pubmed',
                'id': ','.join(batch_ids),
                'retmode': 'xml',
                'api_key': PUBMED_API_KEY
            }
            
            # 添加重试机制
            max_retries = 3
            retry_delay = 1  # 初始延迟1秒
            success = False
            
            for retry in range(max_retries):
                try:
                    # 发送请求获取详情
                    fetch_url = f"{PUBMED_BASE_URL}efetch.fcgi"
                    logger.info(f"获取文献详情 (批次 {current_batch}, 尝试 {retry + 1}/{max_retries}): {fetch_url}")

                    # 使用更稳定的请求配置
                    response = requests.get(
                        fetch_url,
                        params=fetch_params,
                        timeout=(10, 60),  # 连接超时10秒，读取超时60秒
                        headers={
                            'User-Agent': 'NNScholar/1.0 (mailto:support@nnscholar.com)',
                            'Accept': 'application/xml',
                            'Connection': 'keep-alive'
                        },
                        verify=True,
                        allow_redirects=True
                    )
            
                    if response.status_code == 200:
                        # 解析XML响应
                        papers = parse_pubmed_xml(response.content)
                        all_papers.extend(papers)
                        logger.info(f"✓ 第 {current_batch}/{total_batches} 批完成，成功获取 {len(papers)} 篇文献")
                        
                        # 计算当前进度
                        current_count = len(all_papers)
                        percentage = round((current_count / len(id_list)) * 100, 1)
                        logger.info(f"当前进度: {current_count}/{len(id_list)} 篇 ({percentage}%)")
                        
                        # 发送批次完成的进度更新
                        update_fetch_progress(session_id, 'batch_complete', 
                                        f"第 {current_batch}/{total_batches} 批处理完成，已获取 {current_count}/{len(id_list)} 篇文献 ({percentage}%)", 
                                        percentage,
                                        current=current_count, 
                                        total=len(id_list), 
                                        batch_info={'current_batch': current_batch, 
                                                    'total_batches': total_batches, 
                                                    'batch_size': len(batch_ids), 
                                                    'batch_success': len(papers)})
                        success = True
                        break
                    else:
                        logger.warning(f"第 {retry + 1} 次尝试失败: HTTP {response.status_code}")
                        if retry < max_retries - 1:
                            wait_time = min(retry_delay, 5)  # 最多等待5秒
                            logger.info(f"等待 {wait_time} 秒后重试...")
                            time.sleep(wait_time)
                            retry_delay *= 2  # 指数退避

                except (requests.exceptions.ConnectionError,
                        requests.exceptions.Timeout,
                        requests.exceptions.SSLError,
                        requests.exceptions.ChunkedEncodingError) as e:
                    logger.warning(f"第 {retry + 1} 次尝试网络错误: {type(e).__name__}: {str(e)}")
                    if retry < max_retries - 1:
                        wait_time = min(retry_delay, 5)
                        logger.info(f"等待 {wait_time} 秒后重试...")
                        time.sleep(wait_time)
                        retry_delay *= 2
                except Exception as e:
                    logger.error(f"第 {retry + 1} 次尝试出现未预期错误: {str(e)}")
                    if retry < max_retries - 1:
                        time.sleep(1)
                    else:
                        logger.error(f"批次 {current_batch} 处理失败，跳过该批次")
            
            if not success:
                logger.error(f"✗ 第 {current_batch}/{total_batches} 批失败，已重试 {max_retries} 次")
                # 发送批次失败的消息
                update_fetch_progress(session_id, 'batch_error', 
                                   f"第 {current_batch}/{total_batches} 批处理失败，已重试 {max_retries} 次", 
                                   (current_batch / total_batches) * 100,
                                   batch_info={'current_batch': current_batch, 
                                             'total_batches': total_batches, 
                                             'error': f"获取失败，已重试 {max_retries} 次"})
            
            # 添加短暂延时，避免请求过于频繁
            if current_batch < total_batches:  # 最后一批不需要延时
                time.sleep(0.5)
        
        logger.info(f"文献获取完成，共处理 {len(all_papers)}/{len(id_list)} 篇文献")
        
        # 发送完成信息
        final_percentage = round((len(all_papers) / len(id_list)) * 100, 1)
        update_fetch_progress(session_id, 'fetch_complete', 
                            f"文献获取完成，共获取 {len(all_papers)}/{len(id_list)} 篇文献", 
                            final_percentage,
                            current=len(all_papers), 
                            total=len(id_list))
        
        return all_papers
        
    except Exception as e:
        logger.error(f"获取文献详情过程中发生错误: {str(e)}\n{traceback.format_exc()}")
        # 发送错误信息
        if session_id:
            update_fetch_progress(session_id, 'error', 
                                f"获取文献详情失败: {str(e)}", 
                                100,  # 错误时显示100%
                                error=str(e))
        return []

def extract_paper_info(article):
    """从XML中提取文献信息"""
    try:
        # 提取基本信息
        pmid = article.find('PMID').text if article.find('PMID') else None
        if not pmid:
            logger.warning("文献缺少PMID，跳过")
            return None
            
        # logger.info(f"开始处理文献 PMID: {pmid}")
            
        # 提取标题
        title_element = article.find('ArticleTitle')
        title = title_element.text if title_element else None
        if not title:
            logger.warning(f"文献 {pmid} 缺少标题，跳过")
            return None
            
        # logger.info(f"文献标题: {title[:100]}...")
            
        # 提取期刊信息
        journal_element = article.find('Journal')
        if not journal_element:
            logger.warning(f"文献 {pmid} 缺少期刊信息，跳过")
            return None
            
        # 提取期刊标题 - 优先使用Title，如果没有则使用ISOAbbreviation
        journal_title = None
        journal_full = journal_element.find('Title')
        journal_iso = journal_element.find('ISOAbbreviation')
        
        logger.info(f"期刊信息提取详情:")
        logger.info(f"- 完整标题: {journal_full.text if journal_full else 'N/A'}")
        logger.info(f"- ISO缩写: {journal_iso.text if journal_iso else 'N/A'}")
        
        if journal_full and journal_full.text:
            journal_title = journal_full.text
            logger.info(f"使用完整期刊标题: {journal_title}")
        elif journal_iso and journal_iso.text:
            journal_title = journal_iso.text
            logger.info(f"使用期刊ISO缩写: {journal_title}")
            
        if not journal_title:
            logger.warning(f"文献 {pmid} 缺少期刊标题")
            return None
            
        issn = journal_element.find('ISSN').text if journal_element.find('ISSN') else None
        
        logger.info(f"期刊信息 - 标题: {journal_title}, ISSN: {issn}")
        
        if not issn:
            logger.warning(f"文献 {pmid} 缺少ISSN，跳过")
            return None
            
        # 提取发表年份
        pub_year = None
        pub_date = article.find('PubDate')
        
        logger.info(f"开始提取发表年份，PubDate标签内容: {pub_date}")
        
        if pub_date:
            # 尝试从Year标签提取
            year_elem = pub_date.find('Year')
            if year_elem and year_elem.text:
                try:
                    pub_year = int(year_elem.text)
                    logger.info(f"成功提取发表年份: {pub_year}")
                except ValueError:
                    logger.warning(f"无效的年份格式: {year_elem.text}")
            else:
                # 尝试从MedlineDate中提取
                medline_date = pub_date.find('MedlineDate')
                if medline_date and medline_date.text:
                    try:
                        # 提取第一个四位数字作为年份
                        year_match = re.search(r'\b\d{4}\b', medline_date.text)
                        if year_match:
                            pub_year = int(year_match.group())
                            logger.info(f"从MedlineDate提取到年份: {pub_year}")
                    except ValueError:
                        logger.warning(f"无法从MedlineDate提取年份: {medline_date.text}")
        else:
            logger.warning(f"文献 {pmid} 缺少PubDate标签")
        
        if not pub_year:
            logger.warning(f"文献 {pmid} 未能提取到发表年份")
            return None
            
        try:
            pub_year = int(pub_year)
            # logger.info(f"成功将年份转换为整数: {pub_year}")
        except ValueError:
            logger.warning(f"文献 {pmid} 的发表年份格式无效: {pub_year}")
            return None
            
        # 获取期刊指标
        journal_metrics = get_journal_metrics(issn)
        if not journal_metrics:
            logger.warning(f"未找到期刊 {journal_title} (ISSN: {issn}) 的指标信息")
            # 如果没有找到期刊指标，仍然保留期刊基本信息
            journal_metrics = {
                'title': journal_title,
                'issn': issn,
                'impact_factor': 'N/A',
                'jcr_quartile': 'N/A',
                'cas_quartile': 'N/A'
            }
            logger.info("使用默认期刊指标信息")
        else:
            # 确保期刊标题使用从PubMed获取的标题
            journal_metrics['title'] = journal_title
            logger.info(f"获取到期刊指标:")
            logger.info(f"- 期刊标题: {journal_metrics['title']}")
            logger.info(f"- 影响因子: {journal_metrics['impact_factor']}")
            logger.info(f"- JCR分区: {journal_metrics['jcr_quartile']}")
            logger.info(f"- CAS分区: {journal_metrics['cas_quartile']}")
            
        # 构建文献信息
        paper_info = {
            'pmid': pmid,
            'title': title,
            'journal_info': journal_metrics,
            'pub_year': pub_year
        }
        
        logger.info(f"文献 {pmid} 的完整信息:")
        logger.info(f"- 标题: {title}")
        logger.info(f"- 发表年份: {pub_year}")
        logger.info(f"- 期刊标题: {journal_metrics['title']}")
        logger.info(f"- ISSN: {issn}")
        logger.info(f"- 影响因子: {journal_metrics['impact_factor']}")
        logger.info(f"- JCR分区: {journal_metrics['jcr_quartile']}")
        logger.info(f"- CAS分区: {journal_metrics['cas_quartile']}")
        
        # 提取摘要
        abstract_element = article.find('Abstract')
        if abstract_element:
            abstract_text = ' '.join(text.text for text in abstract_element.find_all('AbstractText'))
            paper_info['abstract'] = abstract_text
            logger.info(f"成功提取摘要，长度: {len(abstract_text)}")
        else:
            paper_info['abstract'] = 'No abstract available'
            logger.warning("未找到摘要信息")
            
        # 提取作者信息
        author_list = article.find('AuthorList')
        if author_list:
            authors = []
            for author in author_list.find_all('Author'):
                last_name = author.find('LastName')
                fore_name = author.find('ForeName')
                if last_name and fore_name:
                    authors.append(f"{last_name.text} {fore_name.text}")
            paper_info['authors'] = authors
            logger.info(f"成功提取作者信息，作者数量: {len(authors)}")
            if authors:
                logger.info(f"第一作者: {authors[0]}")
                if len(authors) > 1:
                    logger.info(f"通讯作者: {authors[-1]}")
        else:
            paper_info['authors'] = []
            logger.warning("未找到作者信息")
            
        # 提取DOI
        article_id_list = article.find('ArticleIdList')
        if article_id_list:
            for article_id in article_id_list.find_all('ArticleId'):
                if article_id.get('IdType') == 'doi':
                    paper_info['doi'] = article_id.text
                    logger.info(f"成功提取DOI: {article_id.text}")
                    break
            if 'doi' not in paper_info:
                paper_info['doi'] = ''
                logger.warning(f"文献 {pmid} 未找到DOI信息")
        else:
            paper_info['doi'] = ''
            logger.warning(f"文献 {pmid} 缺少ArticleIdList")
            
        return paper_info
            
    except Exception as e:
        logger.error(f"解析文献信息时出错: {str(e)}\n{traceback.format_exc()}")
        return None

def split_paragraph_to_sentences(paragraph):
    """使用 NLTK 将段落分解为句子"""
    try:
        # 检查是否包含中文字符
        has_chinese = any('\u4e00' <= char <= '\u9fff' for char in paragraph)
        
        if has_chinese:
            # 对于中文文本，使用标点符号分句
            sentences = []
            current_sentence = ""
            # 中文分句标点符号
            end_marks = {'。', '！', '？', '；', '.', '!', '?', ';'}
            
            for char in paragraph:
                current_sentence += char
                if char in end_marks:
                    sentence = current_sentence.strip()
                    if sentence:
                        # 确保句子以标点符号结尾
                        if not any(sentence.endswith(mark) for mark in end_marks):
                            sentence += '。'
                        sentences.append(sentence)
                    current_sentence = ""
            
            # 处理最后一个可能没有结束标点的句子
            if current_sentence.strip():
                sentence = current_sentence.strip()
                # 确保最后一个句子也以标点符号结尾
                if not any(sentence.endswith(mark) for mark in end_marks):
                    sentence += '。'
                sentences.append(sentence)
        else:
            # 对于英文文本，使用NLTK的分句功能
            try:
                sentences = nltk.sent_tokenize(paragraph)
            except LookupError:
                # 如果NLTK数据未下载，使用简单的分句规则
                sentences = [s.strip() for s in re.split('[.!?]+', paragraph) if s.strip()]
        
        # 过滤空句子并去重
        return list(dict.fromkeys(s for s in sentences if s.strip()))
        
    except Exception as e:
        logger.error(f"分句过程中出现错误: {str(e)}")
        # 发生错误时使用最简单的分句方式
        sentences = [s.strip() for s in re.split('[.。!！?？;；]+', paragraph) if s.strip()]
        # 确保每个句子都以标点符号结尾
        return [s if any(s.endswith(mark) for mark in {'。', '！', '？', '；', '.', '!', '?', ';'}) else s + '。' for s in sentences]

async def process_sentence_async(session_id: str, sentence: dict, filters: dict = None) -> dict:
    """异步处理单个句子的检索
    
    Args:
        session_id (str): 用户会话ID
        sentence (dict): 包含文本和检索策略的句子字典
        filters (dict, optional): 筛选条件
        
    Returns:
        dict: 检索结果
    """
    try:
        # 执行检索
        id_list, search_strategy, total_count, filtered_count = search_pubmed(sentence['search_strategy'])
        
        # 获取文献详情
        papers = fetch_paper_details(id_list)
        
        # 应用筛选条件
        if filters:
            filtered_papers, stats = filter_papers_by_metrics(papers, filters)
        else:
            filtered_papers = papers
            stats = {
                'total': len(papers),
                'filtered': len(papers)
            }
        
        return {
            'text': sentence['text'],
            'papers': filtered_papers,
            'search_strategy': sentence['search_strategy'],
            'total_count': total_count,
            'filtered_count': len(filtered_papers)
        }
        
    except Exception as e:
        logger.error(f"处理句子时出错: {str(e)}\n{traceback.format_exc()}")
        return None

async def process_paragraph_async(session_id: str, sentences: List[dict], filters: dict = None) -> List[dict]:
    """异步处理段落中的所有句子
    
    Args:
        session_id (str): 用户会话ID
        sentences (List[dict]): 句子列表
        filters (dict, optional): 筛选条件
        
    Returns:
        List[dict]: 检索结果列表
    """
    try:
        # 创建任务列表
        tasks = [
            process_sentence_async(session_id, sentence, filters)
            for sentence in sentences
        ]
        
        # 并发执行所有任务
        results = await asyncio.gather(*tasks)
        
        # 过滤掉失败的结果
        return [r for r in results if r is not None]
        
    except Exception as e:
        logger.error(f"处理段落时出错: {str(e)}\n{traceback.format_exc()}")
        return []

def process_paragraph_threaded(session_id: str, sentences: List[dict], filters: dict = None) -> List[dict]:
    """使用线程池处理段落中的所有句子
    
    Args:
        session_id (str): 用户会话ID
        sentences (List[dict]): 句子列表
        filters (dict, optional): 筛选条件
        
    Returns:
        List[dict]: 检索结果列表
    """
    try:
        results = []
        with ThreadPoolExecutor(max_workers=len(sentences)) as executor:
            # 提交所有任务
            future_to_sentence = {
                executor.submit(
                    lambda s: process_sentence_async(session_id, s, filters),
                    sentence
                ): sentence
                for sentence in sentences
            }
            
            # 收集结果
            for future in as_completed(future_to_sentence):
                try:
                    result = future.result()
                    if result:
                        results.append(result)
                except Exception as e:
                    logger.error(f"处理句子时出错: {str(e)}")
                    continue
                    
        return results
        
    except Exception as e:
        logger.error(f"处理段落时出错: {str(e)}\n{traceback.format_exc()}")
        return []

def generate_broader_query(query):
    """生成更宽泛的搜索策略"""
    try:
        # 移除一些限制性标签
        broader = query.replace('[Title/Abstract]', '[All Fields]')
        broader = broader.replace('[Mesh]', '[All Fields]')
        
        # 如果包含多个AND条件，只保留部分
        terms = broader.split(' AND ')
        if len(terms) > 2:
            # 保留最重要的2-3个术语
            terms = terms[:3]
            broader = ' AND '.join(terms)
        
        logger.info(f"原始查询: {query}")
        logger.info(f"更宽泛的查询: {broader}")
        return broader
        
    except Exception as e:
        logger.error(f"生成更宽泛查询时出错: {str(e)}")
        return query  # 出错时返回原始查询

def export_papers(papers, query, file_suffix=''):
    """
    将论文信息同时导出为Excel和Word格式
    
    Args:
        papers (list): 论文信息列表
        query (str): 搜索查询
        file_suffix (str): 文件名后缀
    
    Returns:
        tuple: (excel_path, word_path) 导出文件的路径
    """
    try:
        session_id = request.headers.get('sid')
        if not session_id:
            raise ValueError("无效的会话ID")
            
        # 发送导出开始的消息
        update_search_progress(session_id, 'export_progress', "开始导出文献...", 0)
        
        # 如果是段落模式，合并所有子句的文献
        mode = request.args.get('mode', 'single')
        if mode == 'paragraph':
            # 获取所有子句
            sentences = []
            all_papers = []
            
            # 合并所有子句的文献，并记录每个文献对应的句子
            for sentence_result in papers:
                sentence_text = sentence_result.get('text', '')
                sentence_papers = sentence_result.get('papers', [])
                sentences.append(sentence_text)
                
                # 为每篇文献添加对应的句子信息
                for paper in sentence_papers:
                    paper['source_sentence'] = sentence_text
                    all_papers.append(paper)
            
            # 使用合并后的文献列表
            papers = all_papers
            # 更新查询文本为所有句子的组合
            query = "\n".join(sentences)
        
        # 导出Excel
        update_search_progress(session_id, 'export_progress', "正在导出Excel文件...", 30)
        excel_path = export_papers_to_excel(papers, query, file_suffix)
        
        # 发送Excel导出完成的消息
        update_search_progress(session_id, 'export_progress', "Excel文件导出完成", 60)
        
        # 导出Word
        update_search_progress(session_id, 'export_progress', "正在导出Word文件...", 70)
        word_path = export_papers_to_word(papers, query, file_suffix)
        
        # 发送Word导出完成的消息
        update_search_progress(session_id, 'export_progress', "Word文件导出完成", 90)
        
        # 发送导出完成的消息
        update_search_progress(session_id, 'export_progress', "文献导出完成", 100, 
                             files={'excel': excel_path, 'word': word_path})
        
        return excel_path, word_path
        
    except Exception as e:
        logger.error(f'导出文件失败：{str(e)}\n{traceback.format_exc()}')
        
        # 发送错误信息
        if session_id:
            update_search_progress(session_id, 'export_progress', f"导出文件失败: {str(e)}", 100)
            
        return None, None

def export_papers_to_excel(papers, query, file_suffix=''):
    """将文献信息导出为Excel表格"""
    try:
        session_id = request.headers.get('sid')
        if not session_id:
            raise ValueError("无效的会话ID")
            
        # 发送开始导出Excel的消息
        update_search_progress(session_id, 'excel_export_progress', "正在准备数据...", 10)
        
        # 准备数据
        data = []
        for i, paper in enumerate(papers):
            journal_info = paper.get('journal_info', {})
            paper_data = {
                '标题': paper.get('title', ''),
                '摘要': paper.get('abstract', ''),
                '作者': ', '.join(paper.get('authors', [])),
                '发表年份': paper.get('pub_year', ''),
                '期刊名称': journal_info.get('title', ''),
                '影响因子': journal_info.get('impact_factor', 'N/A'),
                'JCR分区': journal_info.get('jcr_quartile', 'N/A'),
                'CAS分区': journal_info.get('cas_quartile', 'N/A'),
                '关键词': ', '.join(paper.get('keywords', [])),
                'DOI': paper.get('doi', ''),
                'PMID': paper.get('pmid', ''),
                '相关度': f"{paper.get('relevance', 0):.1f}%"
            }
            
            # 如果是段落模式，添加来源句子
            if 'source_sentence' in paper:
                paper_data['来源句子'] = paper['source_sentence']
            
            data.append(paper_data)
            
            # 每处理10篇文献发送一次进度更新
            if (i + 1) % 10 == 0:
                progress = min(10 + int((i + 1) / len(papers) * 40), 50)
                update_search_progress(session_id, 'excel_export_progress', 
                                    f"正在处理数据 ({i + 1}/{len(papers)})...", progress)
        
        # 创建DataFrame
        df = pd.DataFrame(data)
        
        # 生成文件名
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_query = re.sub(r'[^\w\u4e00-\u9fff]', '_', query)
        if len(safe_query) > 50:
            safe_query = safe_query[:50]
        filename = f'papers_{safe_query}_{file_suffix}_{timestamp}.xlsx'
        filepath = os.path.join(EXPORTS_DIR, filename)
        
        # 导出到Excel
        df.to_excel(filepath, index=False, engine='openpyxl')
        
        # 发送导出完成的消息
        update_search_progress(session_id, 'excel_export_progress', "Excel文件导出完成", 100)
        
        return filepath
        
    except Exception as e:
        logger.error(f"导出Excel文件时发生错误: {str(e)}")
        if session_id:
            update_search_progress(session_id, 'excel_export_progress', 
                                f"导出Excel文件失败: {str(e)}", 100)
        return None

def export_papers_to_word(papers, query, file_suffix=''):
    """将论文信息导出为Word文档"""
    try:
        session_id = request.headers.get('sid')
        if not session_id:
            raise ValueError("无效的会话ID")
            
        # 发送开始导出Word的消息
        update_search_progress(session_id, 'word_export_progress', "正在创建Word文档...", 10)
        
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        title = doc.add_heading('文献检索报告', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加检索信息
        doc.add_heading('检索策略', level=1)
        doc.add_paragraph(f'检索词：\n{query}')
        doc.add_paragraph(f'检索时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'检索到的文献数量：{len(papers)}')
        
        # 如果是段落模式，添加句子分析部分
        if any('source_sentence' in paper for paper in papers):
            doc.add_heading('句子分析', level=1)
            # 使用字典来收集每个句子的相关文献
            sentence_papers = {}
            for paper in papers:
                if 'source_sentence' in paper:
                    if paper['source_sentence'] not in sentence_papers:
                        sentence_papers[paper['source_sentence']] = []
                    sentence_papers[paper['source_sentence']].append(paper)
            
            # 按句子顺序显示分析
            for i, (sentence, related_papers) in enumerate(sentence_papers.items(), 1):
                # 添加句子标题
                doc.add_heading(f'句子 {i}', level=2)
                
                # 添加原句
                p = doc.add_paragraph()
                p.add_run('原句：').bold = True
                p.add_run(sentence)
                
                # 添加相关文献数量
                doc.add_paragraph(f'相关文献数量：{len(related_papers)}篇')
                
                # 添加该句子的文献列表
                if related_papers:
                    doc.add_heading(f'相关文献列表：', level=3)
                    for j, paper in enumerate(related_papers, 1):
                        # 获取期刊信息
                        journal_info = paper.get('journal_info', {})
                        
                        # 添加文献标题
                        p = doc.add_paragraph()
                        p.add_run(f'{j}. ').bold = True
                        title_run = p.add_run(paper.get('title', 'N/A'))
                        title_run.bold = True
                        
                        # 添加作者信息
                        authors_str = ', '.join(paper.get('authors', [])) if paper.get('authors') else 'N/A'
                        doc.add_paragraph(f'作者：{authors_str}')
                        
                        # 添加期刊信息
                        doc.add_paragraph(f'期刊：{journal_info.get("title", "N/A")}')
                        doc.add_paragraph(f'发表时间：{paper.get("pub_year", "N/A")}')
                        
                        # 添加影响因子和分区信息
                        doc.add_paragraph(f'影响因子：{journal_info.get("impact_factor", "N/A")}')
                        doc.add_paragraph(f'JCR分区：{journal_info.get("jcr_quartile", "N/A")}')
                        doc.add_paragraph(f'CAS分区：{journal_info.get("cas_quartile", "N/A")}')
                        
                        # 添加关键词信息
                        keywords_str = ', '.join(paper.get('keywords', [])) if paper.get('keywords') else 'N/A'
                        doc.add_paragraph(f'关键词：{keywords_str}')
                        
                        # 添加DOI和PMID
                        doc.add_paragraph(f'DOI：{paper.get("doi", "N/A")}')
                        doc.add_paragraph(f'PMID：{paper.get("pmid", "N/A")}')
                        
                        # 添加摘要
                        if paper.get('abstract'):
                            doc.add_paragraph('摘要：').add_run(paper['abstract']).italic = True
                        
                        # 添加分隔线
                        if j < len(related_papers):
                            doc.add_paragraph('_' * 50)
                
                # 在句子之间添加分隔线
                if i < len(sentence_papers):
                    doc.add_paragraph('=' * 50)
        
        # 如果不是段落模式，直接显示文献列表
        else:
            doc.add_heading('文献列表', level=1)
            for i, paper in enumerate(papers, 1):
                # 获取期刊信息
                journal_info = paper.get('journal_info', {})
                
                # 添加文献标题
                p = doc.add_paragraph()
                p.add_run(f'{i}. ').bold = True
                title_run = p.add_run(paper.get('title', 'N/A'))
                title_run.bold = True
                
                # 添加作者信息
                authors_str = ', '.join(paper.get('authors', [])) if paper.get('authors') else 'N/A'
                doc.add_paragraph(f'作者：{authors_str}')
                
                # 添加期刊信息
                doc.add_paragraph(f'期刊：{journal_info.get("title", "N/A")}')
                doc.add_paragraph(f'发表时间：{paper.get("pub_year", "N/A")}')
                
                # 添加影响因子和分区信息
                doc.add_paragraph(f'影响因子：{journal_info.get("impact_factor", "N/A")}')
                doc.add_paragraph(f'JCR分区：{journal_info.get("jcr_quartile", "N/A")}')
                doc.add_paragraph(f'CAS分区：{journal_info.get("cas_quartile", "N/A")}')
                
                # 添加关键词信息
                keywords_str = ', '.join(paper.get('keywords', [])) if paper.get('keywords') else 'N/A'
                doc.add_paragraph(f'关键词：{keywords_str}')
                
                # 添加DOI和PMID
                doc.add_paragraph(f'DOI：{paper.get("doi", "N/A")}')
                doc.add_paragraph(f'PMID：{paper.get("pmid", "N/A")}')
                
                # 添加摘要
                if paper.get('abstract'):
                    doc.add_paragraph('摘要：').add_run(paper['abstract']).italic = True
                
                # 添加分隔线
                if i < len(papers):
                    doc.add_paragraph('_' * 50)
        
        # 生成文件名
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_query = re.sub(r'[^\w\u4e00-\u9fff]', '_', query[:50])  # 限制查询长度为50个字符
        filename = f'papers_{safe_query}_{file_suffix}_{timestamp}.docx'
        filepath = os.path.join(EXPORTS_DIR, filename)
        
        # 保存文档
        doc.save(filepath)
        
        # 发送导出完成的消息
        update_search_progress(session_id, 'word_export_progress', "Word文档导出完成", 100)
        
        return filepath
        
    except Exception as e:
        logger.error(f"导出Word文档时发生错误: {str(e)}")
        if session_id:
            update_search_progress(session_id, 'word_export_progress', 
                                f"导出Word文档失败: {str(e)}", 100)
        return None

@app.route('/')
def index():
    """主页"""
    return render_template('index.html')

@app.route('/chat')
def chat():
    """新版聊天界面"""
    return render_template('chat.html')

@app.route('/test-export')
def test_export():
    """导出功能测试页面"""
    return send_from_directory('.', 'test_export.html')

@app.route('/download/<session_id>/<file_type>')
def download_file(session_id, file_type):
    """动态生成并下载导出文件"""
    try:
        # 验证参数
        if file_type not in ['excel', 'word']:
            return jsonify({'error': '无效的文件类型'}), 400

        # 从缓存中获取文献数据
        with papers_cache_lock:
            if session_id not in papers_cache:
                logger.warning(f"会话 {session_id} 的文献数据不存在")
                return jsonify({'error': '文献数据不存在，请重新检索'}), 404

            papers_data = papers_cache[session_id].copy()

        # 记录下载请求
        logger.info(f"动态生成文件请求: {session_id} - {file_type}")

        # 动态生成文件
        if file_type == 'excel':
            file_content, filename = generate_excel_in_memory(papers_data)
            mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        else:  # word
            file_content, filename = generate_word_in_memory(papers_data)
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        # 创建响应
        response = make_response(file_content)
        response.headers['Content-Type'] = mimetype

        # 处理中文文件名编码问题
        try:
            # 尝试使用ASCII编码文件名
            filename.encode('ascii')
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        except UnicodeEncodeError:
            # 如果包含非ASCII字符，使用RFC 5987编码
            import urllib.parse
            filename_encoded = urllib.parse.quote(filename.encode('utf-8'))
            response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{filename_encoded}"

        response.headers['Content-Length'] = len(file_content)

        logger.info(f"文件生成成功: {filename} ({len(file_content)} bytes)")
        return response

    except Exception as e:
        logger.error(f"动态生成文件时出错: {str(e)}\n{traceback.format_exc()}")
        return jsonify({'error': '文件生成失败'}), 500

@app.route('/api/search', methods=['POST'])
def search():
    """处理搜索请求"""
    try:
        data = request.get_json()
        query = data.get('query', '').strip()
        mode = data.get('mode', 'single')
        filters = data.get('filters', {})

        logger.info(f"接收到的筛选条件: {filters}")
        
        if not query:
            return jsonify({
                'success': False,
                'error': '请输入检索内容'
            }), 400
            
        # 获取当前用户的会话ID
        session_id = request.headers.get('sid')
        if not session_id:
            return jsonify({
                'success': False,
                'error': '无效的会话ID'
            }), 400
            
        # 更新用户活跃状态
        handle_connect(session_id)
        
        # 发送开始搜索的进度信息
        update_search_progress(session_id, 'start', "开始处理搜索请求...", 0)
        
        # 根据模式处理搜索
        if mode == 'paragraph':
            # 段落模式处理
            update_search_progress(session_id, 'paragraph_mode', "正在使用段落模式处理...", 10)
            
            # 分解段落为句子
            sentences = split_paragraph_to_sentences(query)
            if not sentences:
                return jsonify({
                    'success': True,
                    'sentences': []
                })
            
            # 使用线程池并行处理所有句子
            results = process_paragraph_threaded(session_id, sentences, filters)
            
            # 发送完成消息
            update_search_progress(session_id, 'complete', f"段落处理完成，共处理 {len(results)} 个句子", 100)
            
            return jsonify({
                'success': True,
                'sentences': results
            })
            
        elif mode == 'strategy':
            # 直接使用提供的检索策略
            search_strategy = query
            logger.info(f"使用提供的检索策略: {search_strategy}")
            
            update_search_progress(session_id, 'executing_strategy', "正在执行检索策略...", 30)
        else:
            # 生成检索策略
            update_search_progress(session_id, 'generating_strategy', "正在生成检索策略...", 20)
            
            prompt = """作为PubMed搜索专家，请为以下研究内容生成优化的PubMed检索策略：

研究内容：{query}

要求：
1. 提取2-3个核心概念，每个概念扩展：
   - 首选缩写（如有）
   - 全称术语
   - 相近术语和同义词
   - 仅返回检索策略，不要其他解释

2. 结构要求：
   (
     ("缩写"[Title/Abstract] OR "全称术语"[Title/Abstract] OR "同义词"[Title/Abstract])
     AND
     ("缩写"[Title/Abstract] OR "全称术语"[Title/Abstract] OR "同义词"[Title/Abstract])
   )

3. 强制规则：
   - 每个概念最多3个术语（缩写+全称+同义词）
   - 只使用Title/Abstract字段，不使用MeSH
   - 保持AND连接的逻辑组不超过3组
   - 使用精确匹配，所有术语都要加双引号"""

            try:
                search_strategy = call_deepseek_api(prompt.format(query=query))
                logger.info(f"生成的检索策略: {search_strategy}")
                
                update_search_progress(session_id, 'strategy_generated', "检索策略生成完成", 40, search_strategy=search_strategy)
            except Exception as e:
                logger.warning(f"DeepSeek API调用失败，使用基本搜索策略: {str(e)}")
                # 使用更宽松的基本搜索策略
                search_strategy = f'({query}[Title/Abstract] OR {query}[All Fields])'
                logger.info(f"使用基本搜索策略: {search_strategy}")

                update_search_progress(session_id, 'strategy_fallback', "检索策略生成失败，使用基本策略", 40, search_strategy=search_strategy)
        
        # 执行搜索
        id_list, _, total_count, _ = search_pubmed(search_strategy)

        # 如果没有找到文献，尝试更简单的搜索策略
        if not id_list:
            logger.warning(f"使用原始策略未找到文献，尝试简化策略")
            simple_strategy = f'{query}[All Fields]'
            logger.info(f"尝试简化策略: {simple_strategy}")
            update_search_progress(session_id, 'retry_search', "尝试简化搜索策略...", 50)
            id_list, _, total_count, _ = search_pubmed(simple_strategy)
            search_strategy = simple_strategy
        
        # 获取文献详情
        update_search_progress(session_id, 'fetching_details', "正在获取文献详情...", 60)
        
        papers = fetch_paper_details(id_list)
        
        # 计算相关性得分
        update_search_progress(session_id, 'calculating_relevance', "正在计算相关性得分...", 80)

        # 使用嵌入模型批量计算相关度
        try:
            logger.info(f"开始批量计算 {len(papers)} 篇文献的嵌入相关度")
            batch_relevance_scores = calculate_batch_embedding_relevance(query, papers, EMBEDDING_BATCH_SIZE)

            if batch_relevance_scores:
                # 使用批量计算的结果
                for i, paper in enumerate(papers):
                    if i in batch_relevance_scores:
                        paper['relevance'] = batch_relevance_scores[i]
                    else:
                        # 如果批量计算中没有这篇文献，使用单个嵌入计算
                        paper['relevance'] = calculate_relevance_improved(query, paper)
                logger.info(f"批量嵌入相关度计算成功，处理了 {len(batch_relevance_scores)} 篇文献")
            else:
                # 如果批量计算失败，回退到逐个嵌入计算
                logger.warning("批量嵌入相关度计算失败，回退到逐个嵌入计算")
                for paper in papers:
                    paper['relevance'] = calculate_relevance_improved(query, paper)
        except Exception as e:
            logger.error(f"嵌入相关度计算出错: {str(e)}，使用默认分数")
            # 如果嵌入计算完全失败，给所有文献一个基础分数
            for paper in papers:
                paper['relevance'] = 50.0  # 给一个中等分数
        
        # 应用筛选条件
        filtered_papers, stats = filter_papers_by_metrics(papers, filters)
        
        # 保存文献数据到内存缓存
        update_search_progress(session_id, 'caching', "正在缓存文献数据...", 90)

        # 将文献数据保存到内存缓存中
        with papers_cache_lock:
            papers_cache[session_id] = {
                'papers': filtered_papers,
                'original_papers': papers,
                'query': query,
                'search_strategy': search_strategy,
                'total_count': total_count,
                'filtered_count': len(filtered_papers),
                'timestamp': datetime.now(),
                'stats': stats
            }

        # 发送完成消息
        update_search_progress(session_id, 'complete', f"搜索完成，找到 {len(filtered_papers)} 篇相关文献", 100)

        return jsonify({
            'success': True,
            'data': filtered_papers,
            'original_papers': papers,
            'search_strategy': search_strategy,
            'total_count': total_count,
            'filtered_count': len(filtered_papers),
            'has_cached_data': True,  # 标识有缓存数据可供导出
            'session_id': session_id  # 返回会话ID用于下载
        })
        
    except Exception as e:
        logger.error(f"搜索请求处理出错: {str(e)}\n{traceback.format_exc()}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/messages', methods=['GET'])
def get_messages():
    """获取用户消息历史"""
    try:
        session_id = request.headers.get('sid')
        if not session_id:
            return jsonify({
                'success': False,
                'error': '无效的会话ID'
            }), 400
            
        with user_stats_lock:
            messages = get_user_messages(session_id)
            return jsonify({
            'success': True,
            'messages': messages
        })
        
    except Exception as e:
        logger.error(f"获取消息历史出错: {str(e)}\n{traceback.format_exc()}")
        return jsonify({
            'success': False,
            'error': '服务器内部错误'
        }), 500

@app.route('/api/session', methods=['DELETE'])
def clear_session():
    """清除用户会话"""
    try:
        session_id = request.headers.get('sid')
        if not session_id:
            return jsonify({
                'success': False,
                'error': '无效的会话ID'
            }), 400
        
        handle_disconnect(session_id)
        return jsonify({
            'success': True,
            'message': '会话已清除'
        })
    
    except Exception as e:
        logger.error(f"清除会话出错: {str(e)}\n{traceback.format_exc()}")
        return jsonify({
            'success': False,
            'error': '服务器内部错误'
        }), 500

@app.route('/api/stats')
def get_stats():
    """获取用户统计信息"""
    try:
        with user_stats_lock, performance_stats_lock:
            current_stats = {
                'total_visits': user_stats['total_visits'],
                'concurrent_users': user_stats['concurrent_users'],
                'peak_concurrent_users': user_stats['peak_concurrent_users'],
                'hourly_visits': get_hourly_visits(),
                'performance': {
                    'avg_response_time': performance_stats['avg_response_time'],
                    'error_rate': (performance_stats['error_count'] / performance_stats['request_count'] * 100) if performance_stats['request_count'] > 0 else 0
                }
            }
        return jsonify(current_stats)
    except Exception as e:
        logger.error(f"获取统计信息时出错: {str(e)}")
        return jsonify({
            'error': '获取统计信息失败'
        }), 500

@app.route('/admin')
def admin():
    """后台管理页面"""
    return render_template('admin.html')

# 启动所有后台线程
def start_background_tasks():
    """启动所有后台任务"""
    try:
        # 启动用户清理线程
        cleaning_thread = threading.Thread(target=clean_inactive_users, daemon=True)
        cleaning_thread.start()
        
        # 启动性能监控线程
        monitor_thread = threading.Thread(target=monitor_system_performance, daemon=True)
        monitor_thread.start()
        
        # 启动导出文件清理线程
        exports_cleaning_thread = threading.Thread(target=clean_exports_directory, daemon=True)
        exports_cleaning_thread.start()
        
        logger.info("所有后台任务已启动")
    except Exception as e:
        logger.error(f"启动后台任务时出错: {str(e)}")

def clean_inactive_users():
    """
    清理不活跃用户的后台任务
    
    每隔一定时间检查并清理超过MAX_CACHE_TIME（24小时）未活跃的用户会话
    """
    logger.info("启动用户清理线程")
    while True:
        try:
            current_time = datetime.now()
            inactive_users = []
            
            # 使用锁保护对用户数据的访问
            with user_stats_lock:
                # 检查所有用户
                for session_id, user_data in active_users.items():
                    last_active = user_data['last_active']
                    # 计算不活跃时间（秒）
                    inactive_time = (current_time - last_active).total_seconds()
                    
                    # 如果超过最大缓存时间，加入待清理列表
                    if inactive_time > MAX_CACHE_TIME:
                        inactive_users.append(session_id)
                
                # 清理不活跃用户
                for session_id in inactive_users:
                    logger.info(f"清理不活跃用户: {session_id}")
                    handle_disconnect(session_id)
                
                if inactive_users:
                    logger.info(f"已清理 {len(inactive_users)} 个不活跃用户")
                else:
                    logger.debug("没有需要清理的不活跃用户")
            
            # 每小时检查一次
            time.sleep(3600)
            
        except Exception as e:
            logger.error(f"清理不活跃用户时出错: {str(e)}\n{traceback.format_exc()}")
            # 发生错误时等待一段时间后继续
            time.sleep(300)

def generate_excel_in_memory(papers_data):
    """在内存中生成Excel文件"""
    try:
        papers = papers_data.get('papers', [])
        query = papers_data.get('query', '未知查询')

        # 创建DataFrame
        df_data = []
        for i, paper in enumerate(papers, 1):
            # 处理JCR分区
            jcr_quartile = paper.get('journal_info', {}).get('jcr_quartile', '')
            jcr_display = f"Q{jcr_quartile}" if jcr_quartile else ''

            # 处理中科院分区
            cas_quartile = paper.get('journal_info', {}).get('cas_quartile', '')
            cas_display = f"{cas_quartile}区" if cas_quartile else ''

            df_data.append({
                '序号': i,
                '标题': paper.get('title', ''),
                '作者': ', '.join(paper.get('authors', [])),
                '期刊': paper.get('journal_info', {}).get('title', ''),
                '发表年份': paper.get('pub_year', ''),
                '影响因子': paper.get('journal_info', {}).get('impact_factor', ''),
                'JCR分区': jcr_display,
                '中科院分区': cas_display,
                '相关度': f"{paper.get('relevance', 0):.1f}%",
                'PMID': paper.get('pmid', ''),
                'DOI': paper.get('doi', ''),
                '摘要': paper.get('abstract', '')
            })

        df = pd.DataFrame(df_data)

        # 使用BytesIO创建内存中的Excel文件
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='文献数据', index=False)

            # 获取工作表并设置格式
            worksheet = writer.sheets['文献数据']

            # 设置列宽
            column_widths = {
                'A': 8,   # 序号
                'B': 50,  # 标题
                'C': 30,  # 作者
                'D': 25,  # 期刊
                'E': 12,  # 年份
                'F': 12,  # 影响因子
                'G': 12,  # JCR分区
                'H': 12,  # 中科院分区
                'I': 12,  # 相关度
                'J': 15,  # PMID
                'K': 20,  # DOI
                'L': 80   # 摘要
            }

            for col, width in column_widths.items():
                worksheet.column_dimensions[col].width = width

        output.seek(0)
        file_content = output.getvalue()

        # 生成文件名（使用英文避免编码问题）
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # 清理查询字符串，移除所有特殊字符和换行符
        clean_query = re.sub(r'[^\w]', '_', query.replace('\n', ' ').replace('\r', ' ').strip())[:20]
        # 确保文件名不包含连续的下划线
        clean_query = re.sub(r'_+', '_', clean_query).strip('_')
        # 如果清理后为空，使用默认名称
        if not clean_query:
            clean_query = "literature_search"
        filename = f"papers_data_{clean_query}_{timestamp}.xlsx"

        return file_content, filename

    except Exception as e:
        logger.error(f"生成Excel文件失败: {str(e)}")
        raise

def generate_word_in_memory(papers_data):
    """在内存中生成Word文件"""
    try:
        papers = papers_data.get('papers', [])
        query = papers_data.get('query', '未知查询')

        # 创建Word文档
        doc = Document()

        # 添加标题
        title = doc.add_heading(f'文献检索报告 - {query}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加基本信息
        doc.add_heading('检索信息', level=1)
        info_para = doc.add_paragraph()
        info_para.add_run('检索主题: ').bold = True
        info_para.add_run(query)
        info_para.add_run('\n检索时间: ').bold = True
        info_para.add_run(datetime.now().strftime("%Y年%m月%d日 %H:%M:%S"))
        info_para.add_run('\n文献数量: ').bold = True
        info_para.add_run(str(len(papers)))

        # 添加文献列表
        doc.add_heading('文献列表', level=1)

        for i, paper in enumerate(papers, 1):
            # 文献标题
            title_para = doc.add_paragraph()
            title_para.add_run(f'{i}. ').bold = True
            title_para.add_run(paper.get('title', '无标题')).bold = True

            # 文献信息
            info_para = doc.add_paragraph()
            authors = ', '.join(paper.get('authors', []))
            if authors:
                info_para.add_run('作者: ').bold = True
                info_para.add_run(authors + '\n')

            journal = paper.get('journal_info', {}).get('title', '')
            if journal:
                info_para.add_run('期刊: ').bold = True
                info_para.add_run(journal + '\n')

            # 影响因子
            impact_factor = paper.get('journal_info', {}).get('impact_factor', '')
            if impact_factor:
                info_para.add_run('影响因子: ').bold = True
                info_para.add_run(str(impact_factor) + '\n')

            # JCR分区
            jcr_quartile = paper.get('journal_info', {}).get('jcr_quartile', '')
            if jcr_quartile:
                info_para.add_run('JCR分区: ').bold = True
                info_para.add_run(f'Q{jcr_quartile}\n')

            # 中科院分区
            cas_quartile = paper.get('journal_info', {}).get('cas_quartile', '')
            if cas_quartile:
                info_para.add_run('中科院分区: ').bold = True
                info_para.add_run(f'{cas_quartile}区\n')

            pub_year = paper.get('pub_year', '')
            if pub_year:
                info_para.add_run('年份: ').bold = True
                info_para.add_run(str(pub_year) + '\n')

            relevance = paper.get('relevance', 0)
            info_para.add_run('相关度: ').bold = True
            info_para.add_run(f'{relevance:.1f}%\n')

            pmid = paper.get('pmid', '')
            if pmid:
                info_para.add_run('PMID: ').bold = True
                info_para.add_run(pmid + '\n')

            # 摘要
            abstract = paper.get('abstract', '')
            if abstract:
                abstract_para = doc.add_paragraph()
                abstract_para.add_run('摘要: ').bold = True
                abstract_para.add_run(abstract)

            # 添加分隔线
            if i < len(papers):
                doc.add_paragraph('─' * 50)

        # 保存到内存
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        file_content = output.getvalue()

        # 生成文件名（使用英文避免编码问题）
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # 清理查询字符串，移除所有特殊字符和换行符
        clean_query = re.sub(r'[^\w]', '_', query.replace('\n', ' ').replace('\r', ' ').strip())[:20]
        # 确保文件名不包含连续的下划线
        clean_query = re.sub(r'_+', '_', clean_query).strip('_')
        # 如果清理后为空，使用默认名称
        if not clean_query:
            clean_query = "literature_search"
        filename = f"research_report_{clean_query}_{timestamp}.docx"

        return file_content, filename

    except Exception as e:
        logger.error(f"生成Word文件失败: {str(e)}")
        raise

def clean_exports_directory():
    """
    清理导出目录中的旧文件

    每隔一定时间检查并清理超过MAX_CACHE_TIME（24小时）的导出文件
    """
    logger.info("启动导出文件清理线程")
    while True:
        try:
            current_time = datetime.now()
            files_removed = 0
            
            # 遍历导出目录中的所有文件
            for file_path in Path(EXPORTS_DIR).glob('*.*'):
                try:
                    # 获取文件的修改时间
                    file_mtime = datetime.fromtimestamp(file_path.stat().st_mtime)
                    # 计算文件存在时间（秒）
                    file_age = (current_time - file_mtime).total_seconds()
                    
                    # 如果文件超过最大缓存时间，删除它
                    if file_age > MAX_CACHE_TIME:
                        file_path.unlink()
                        files_removed += 1
                        logger.debug(f"已删除过期文件: {file_path}")
                except Exception as e:
                    logger.warning(f"删除文件 {file_path} 时出错: {str(e)}")
            
            if files_removed > 0:
                logger.info(f"已清理 {files_removed} 个过期导出文件")
            else:
                logger.debug("没有需要清理的过期导出文件")
            
            # 每6小时检查一次
            time.sleep(21600)
            
        except Exception as e:
            logger.error(f"清理导出文件时出错: {str(e)}\n{traceback.format_exc()}")
            # 发生错误时等待一段时间后继续
            time.sleep(300)

@app.route('/api/generate_strategy', methods=['POST'])
def generate_strategy():
    """生成检索策略"""
    try:
        data = request.get_json()
        query = data.get('query', '').strip()
        mode = data.get('mode', 'single')
        filters = data.get('filters', {})
        
        if not query:
            return jsonify({
                'success': False,
                'error': '请输入检索内容'
            }), 400
            
        # 获取当前用户的会话ID
        session_id = request.headers.get('sid')
        if not session_id:
            return jsonify({
                'success': False,
                'error': '无效的会话ID'
            }), 400
            
        # 更新用户活跃状态
        handle_connect(session_id)
        
        # 发送开始生成的进度信息
        update_search_progress(session_id, 'strategy_progress', "开始生成检索策略...", 0)
        
        if mode == 'paragraph':
            # 段落模式：分句并为每个句子生成检索策略
            sentences = split_paragraph_to_sentences(query)
            if not sentences:
                return jsonify({
                    'success': True,
                    'sentences': []
                })
            
            # 为每个句子生成检索策略
            results = []
            for i, sentence in enumerate(sentences, 1):
                try:
                    # 更新进度
                    progress = (i / len(sentences)) * 100
                    update_search_progress(session_id, 'strategy_progress', 
                                        f"正在为第 {i}/{len(sentences)} 个句子生成检索策略", progress)
                    
                    # 生成检索策略
                    prompt = f"""作为PubMed搜索专家，请为以下研究内容生成优化的PubMed检索策略：

研究内容：{sentence}

要求：
1. 提取2-3个核心概念，每个概念扩展：
   - 首选缩写（如有）
   - 全称术语
   - 相近术语和同义词
   - 仅返回检索策略，不要其他解释

2. 结构要求：
   (
     ("缩写"[Title/Abstract] OR "全称术语"[Title/Abstract] OR "同义词"[Title/Abstract]) 
     AND 
     ("缩写"[Title/Abstract] OR "全称术语"[Title/Abstract] OR "同义词"[Title/Abstract])
   )

3. 强制规则：
   - 每个概念最多3个术语（缩写+全称+同义词）
   - 只使用Title/Abstract字段，不使用MeSH
   - 保持AND连接的逻辑组不超过3组
   - 使用精确匹配，所有术语都要加双引号"""

                    search_strategy = call_deepseek_api(prompt)
                    logger.info(f"为句子生成的检索策略: {search_strategy}")
                    
                    # 构建年份限制
                    year_filter = ""
                    if filters and 'year_start' in filters and 'year_end' in filters:
                        year_start = filters.get('year_start')
                        year_end = filters.get('year_end')
                        if year_start and year_end:
                            year_filter = f' AND ("{year_start}"[Date - Publication] : "{year_end}"[Date - Publication])'
                    
                    # 添加年份限制
                    if year_filter:
                        search_strategy = f"({search_strategy}){year_filter}"
                    
                    results.append({
                        'text': sentence,
                        'search_strategy': search_strategy
                    })
                    
                    # 添加日志
                    logger.info(f"第 {i} 个句子的检索策略生成完成")
                    
                except Exception as e:
                    logger.error(f"处理句子时出错: {str(e)}")
                    continue
            
            # 发送完成消息
            update_search_progress(session_id, 'strategy_complete', 
                                f"检索策略生成完成，共处理 {len(results)} 个句子", 100)
            
            return jsonify({
                'success': True,
                'sentences': results
            })
            
        else:
            # 单句模式：生成单个检索策略
            prompt = f"""作为PubMed搜索专家，请为以下研究内容生成优化的PubMed检索策略：

研究内容：{query}

要求：
1. 提取2-3个核心概念，每个概念扩展：
   - 首选缩写（如有）
   - 全称术语
   - 相近术语和同义词
   - 仅返回检索策略，不要其他解释

2. 结构要求：
   (
     ("缩写"[Title/Abstract] OR "全称术语"[Title/Abstract] OR "同义词"[Title/Abstract]) 
     AND 
     ("缩写"[Title/Abstract] OR "全称术语"[Title/Abstract] OR "同义词"[Title/Abstract])
   )

3. 强制规则：
   - 每个概念最多3个术语（缩写+全称+同义词）
   - 只使用Title/Abstract字段，不使用MeSH
   - 保持AND连接的逻辑组不超过3组
   - 使用精确匹配，所有术语都要加双引号"""

            try:
                search_strategy = call_deepseek_api(prompt)
                logger.info(f"生成的检索策略: {search_strategy}")
                
                # 构建年份限制
                year_filter = ""
                if filters and 'year_start' in filters and 'year_end' in filters:
                    year_start = filters.get('year_start')
                    year_end = filters.get('year_end')
                    if year_start and year_end:
                        year_filter = f' AND ("{year_start}"[Date - Publication] : "{year_end}"[Date - Publication])'
                
                # 添加年份限制
                if year_filter:
                    search_strategy = f"({search_strategy}){year_filter}"
                
                update_search_progress(session_id, 'strategy_complete', "检索策略生成完成", 100)
                
                return jsonify({
                    'success': True,
                    'search_strategy': search_strategy
                })
                
            except Exception as e:
                logger.error(f"生成检索策略时出错: {str(e)}")
                update_search_progress(session_id, 'strategy_error', f"生成检索策略失败: {str(e)}", 100)
                return jsonify({
                    'success': False,
                    'error': str(e)
                }), 500
                
    except Exception as e:
        logger.error(f"生成检索策略时出错: {str(e)}\n{traceback.format_exc()}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500



@app.route('/api/analyze-journal', methods=['POST'])
def analyze_journal():
    """期刊分析API端点"""
    try:
        data = request.get_json()
        journal = data.get('journal')
        keywords = data.get('keywords', '').strip()  # 可选参数
        start_year = data.get('start_year')
        end_year = data.get('end_year')
        
        # 参数验证
        if not journal:
            return jsonify({
                'success': False,
                'error': '请输入期刊名称'
            }), 400
            
        if not all([start_year, end_year]):
            return jsonify({
                'success': False,
                'error': '请输入完整的时间范围'
            }), 400
            
        try:
            start_year = int(start_year)
            end_year = int(end_year)
            if start_year > end_year:
                return jsonify({
                    'success': False,
                    'error': '起始年份不能大于结束年份'
                }), 400
        except ValueError:
            return jsonify({
                'success': False,
                'error': '年份格式无效'
            }), 400
            
        # 创建分析器实例
        analyzer = JournalAnalyzer()
        
        # 构建基本检索策略
        base_query = f"{journal}[ta] AND ({start_year}[pdat]:{end_year}[pdat])"
        
        # 如果有关键词，扩展关键词并添加到检索策略
        if keywords:
            expanded_keywords = expand_keywords(keywords)
            if expanded_keywords:
                base_query = f"({base_query}) AND ({expanded_keywords})"
        
        # 获取文章数据
        articles = analyzer.fetch_journal_articles(base_query)
        
        if not articles:
            return jsonify({
                'success': False,
                'error': '未找到相关文章'
            }), 404
            
        # 分析热点主题
        hot_topics = analyzer.analyze_hot_topics(articles)
        
        # 生成热力图数据 - 转换为简单的[topic, count]格式
        heatmap_data = []
        for topic in hot_topics[:30]:  # 只取前30个主题
            heatmap_data.append([topic['topic'], topic['article_count']])
        
        # 生成词云数据
        wordcloud_data = analyzer.extract_keywords(articles)[:50]  # 限制50个关键词
        
        # 生成趋势数据
        trend_data = analyzer.analyze_trends(articles)
        
        # 分析热点作者
        hot_authors = analyzer.analyze_hot_authors(articles)
        
        # 准备返回数据
        response_data = {
            'success': True,
            'heatmap_data': heatmap_data,
            'wordcloud_data': wordcloud_data,
            'trend_data': trend_data,
            'hot_authors': hot_authors,
            'total_articles': len(articles)
        }
        
        logger.info(f"分析完成，数据预览:")
        logger.info(f"热力图数据: {len(heatmap_data)} 个主题")
        logger.info(f"词云数据: {len(wordcloud_data)} 个关键词")
        logger.info(f"趋势数据: {len(trend_data.get('topics', []))} 个主题, {len(trend_data.get('years', []))} 年")
        logger.info(f"热点作者: {len(hot_authors)} 位")
        
        return jsonify(response_data)
        
    except Exception as e:
        logger.error(f"期刊分析失败: {str(e)}\n{traceback.format_exc()}")
        return jsonify({
            'success': False,
            'error': f'分析失败: {str(e)}'
        }), 500

def expand_keywords(keywords):
    """
    扩展关键词，只处理缩写和全称的转换
    
    Args:
        keywords (str): 用逗号分隔的关键词字符串
        
    Returns:
        str: 扩展后的PubMed检索策略
    """
    if not keywords:
        return ""
        
    # 分割关键词
    keyword_list = [k.strip() for k in keywords.split(',')]
    expanded_terms = []
    
    for keyword in keyword_list:
        if not keyword:
            continue
            
        # 调用DeepSeek API进行扩展
        prompt = f"""作为PubMed检索专家，请为以下关键词生成检索策略，只考虑缩写和全称的转换，不要添加额外相关概念：

关键词：{keyword}

要求：
1. 只扩展缩写和全称的对应关系，例如：
   - "LLM" -> ("LLM"[Title/Abstract] OR "Large Language Model"[Title/Abstract])
   - "CT" -> ("CT"[Title/Abstract] OR "Computed Tomography"[Title/Abstract])
2. 不要添加其他相关概念或同义词
3. 使用Title/Abstract字段
4. 所有术语都要加双引号
5. 直接返回检索策略，不要其他解释"""

        try:
            expanded = call_deepseek_api(prompt)
            expanded_terms.append(expanded)
        except Exception as e:
            logger.warning(f"扩展关键词 {keyword} 时出错: {str(e)}")
            # 如果扩展失败，使用原始关键词
            expanded_terms.append(f'"{keyword}"[Title/Abstract]')
    
    # 将所有扩展后的词组用 AND 连接
    if expanded_terms:
        return " AND ".join(expanded_terms)
    return ""

@app.route('/api/analyze_research_status', methods=['POST'])
def analyze_research_status():
    """分析研究现状API端点"""
    try:
        data = request.get_json()
        query = data.get('query', '').strip()
        titles = data.get('titles', [])
        paper_count = data.get('paper_count', 0)

        if not query:
            return jsonify({
                'success': False,
                'error': '请提供研究方向'
            }), 400

        # 如果没有提供titles，从缓存中获取
        if not titles:
            session_id = request.headers.get('sid') or session.get('session_id')
            logger.info(f"研究现状分析 - 会话ID: {session_id}")
            logger.info(f"研究现状分析 - 缓存中的会话: {list(papers_cache.keys())}")
            if session_id:
                with papers_cache_lock:
                    if session_id in papers_cache:
                        papers_data = papers_cache[session_id]
                        titles = [paper.get('title', '') for paper in papers_data.get('papers', [])]
                        query = papers_data.get('query', query)  # 使用缓存中的查询词
                        logger.info(f"研究现状分析 - 从缓存获取到 {len(titles)} 篇文献")
                    else:
                        logger.warning(f"研究现状分析 - 会话ID {session_id} 不在缓存中")

        if not titles:
            return jsonify({
                'success': False,
                'error': '没有文献标题用于分析，请先进行文献检索'
            }), 400

        # 构建分析提示词 - 使用相关性最强的150篇文献
        selected_titles = titles[:150]  # 取前150篇相关性最强的文献
        titles_text = '\n'.join([f"{i+1}. {title}" for i, title in enumerate(selected_titles)])

        prompt = f"""- Role: 学术研究前沿分析师
- Background: 用户已经收集了特定研究领域的文献资料，需要对这些文献进行深入分析，以生成一份能够反映该领域前沿发展的调研报告。这表明用户对该领域有一定的了解，但需要专业的分析来梳理研究进展、挖掘潜在趋势，并提供有价值的见解。
- Profile: 你是一位在学术研究领域有着丰富经验的前沿分析师，擅长从大量文献中提取关键信息，识别研究趋势和创新点，能够运用专业的分析方法和工具，为不同领域的研究者提供深入且具有前瞻性的调研报告。
- Skills: 你具备文献综述能力、数据分析能力、趋势预测能力以及报告撰写能力。能够快速理解不同领域的核心概念，运用批判性思维评估研究质量，识别文献中的空白和潜在的研究方向。
- Goals:
  1. 对文献进行全面梳理，提取关键研究内容和成果。
  2. 分析该领域的研究趋势，识别当前的研究热点和未来的发展方向。
  3. 挖掘文献中的潜在问题和未解决的挑战，为后续研究提供参考。
  4. 撰写一份结构清晰、内容详实的前沿发展调研报告，包括研究背景、现状、趋势、挑战和建议等部分。
- Constrains: 报告应基于文献内容，确保信息的准确性和客观性。避免引入无关内容，保持报告的聚焦性和专业性。同时，报告应具有一定的深度和广度，能够为该领域的研究者提供有价值的参考。在分析时，请直接引用和讨论文献内容，不要说"基于提供的文献标题"等表述，而是自然地说"文献显示"、"研究表明"、"相关文献"等。
- OutputFormat: 请用中文回答，使用HTML格式，包含适当的标题和段落标签。报告应包括以下部分：
  - 研究现状：详细描述当前的研究进展，包括主要研究成果、研究方法和技术手段等。
  - 研究趋势：分析该领域的研究趋势，包括热点问题、新兴技术、跨学科研究等。
  - 面临的挑战：指出该领域目前存在的主要问题和挑战，如技术瓶颈、数据不足、理论争议等。
  - 未来发展方向：基于现状和趋势，预测该领域的未来发展方向，提出可能的研究方向和建议。

现在请对"{query}"领域进行专业的前沿发展分析。以下是{len(selected_titles)}篇高相关性文献：

{titles_text}

请严格按照上述要求进行分析，在回复中自然地引用文献内容，使用"文献显示"、"研究表明"、"相关文献"等表述，确保内容专业、准确、具有前瞻性。"""

        try:
            # 调用DeepSeek API
            analysis_result = call_deepseek_api(prompt)

            return jsonify({
                'success': True,
                'analysis': analysis_result,
                'analyzed_papers': len(titles),
                'total_papers': paper_count
            })

        except Exception as e:
            logger.error(f"DeepSeek API调用失败: {str(e)}")
            # 提供备用分析（符合学术研究前沿分析师角色）
            fallback_analysis = f"""
            <h3>📊 {query} - 前沿发展分析</h3>

            <h4>🔬 研究现状</h4>
            <p>"{query}"领域当前呈现出多元化发展态势。研究内容涵盖基础理论探索、技术方法创新和临床应用转化等多个层面，显示出该领域的活跃度和重要性。从{len(selected_titles)}篇高相关性文献可以看出，该领域研究深度和广度都在不断扩展。</p>

            <h4>📈 研究趋势</h4>
            <p>该领域正朝着精准化、个性化和智能化方向发展。跨学科融合趋势明显，新兴技术的应用日益广泛，为传统研究范式带来了新的突破。文献显示出明显的技术驱动和临床需求导向特征。</p>

            <h4>⚠️ 面临的挑战</h4>
            <p>当前该领域面临的主要挑战包括：研究方法的标准化需求、数据质量和可重复性问题、以及从实验室研究向临床实践转化的复杂性。这些挑战需要学术界和产业界的共同努力来解决。</p>

            <h4>🚀 未来发展方向</h4>
            <p>未来该领域有望在技术创新、临床转化和产业应用方面取得重大突破。建议加强国际合作、完善研究标准、推动产学研一体化发展，以实现更大的社会价值和科学意义。</p>

            <p><em>💡 注：由于AI分析服务暂时不可用，以上为专业基础分析。</em></p>
            """

            return jsonify({
                'success': True,
                'analysis': fallback_analysis,
                'analyzed_papers': len(titles),
                'total_papers': paper_count
            })

    except Exception as e:
        logger.error(f"研究现状分析出错: {str(e)}\n{traceback.format_exc()}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/review_topic_suggestion', methods=['POST'])
def review_topic_suggestion():
    """基于真实文献的综述选题建议API端点"""
    try:
        data = request.get_json()
        query = data.get('query', '').strip()
        titles = data.get('titles', [])
        papers = data.get('papers', [])

        if not query:
            return jsonify({
                'success': False,
                'error': '请提供研究方向'
            }), 400

        # 如果没有提供titles和papers，从缓存中获取
        if not titles and not papers:
            session_id = request.headers.get('sid') or session.get('session_id')
            if session_id:
                with papers_cache_lock:
                    if session_id in papers_cache:
                        papers_data = papers_cache[session_id]
                        papers = papers_data.get('papers', [])
                        titles = [paper.get('title', '') for paper in papers]
                        query = papers_data.get('query', query)  # 使用缓存中的查询词

        if not titles and not papers:
            return jsonify({
                'success': False,
                'error': '没有文献数据用于分析，请先进行文献检索'
            }), 400

        # 使用前150篇相关性最强的文献
        selected_titles = titles[:150] if titles else []
        selected_papers = papers[:150] if papers else []

        # 构建综述选题分析提示词
        titles_text = '\n'.join([f"{i+1}. {title}" for i, title in enumerate(selected_titles)])

        # 如果有完整的文献信息，提取更多细节
        papers_info = ""
        if selected_papers:
            papers_info = "\n\n文献详细信息（前20篇）：\n"
            for i, paper in enumerate(selected_papers[:20]):
                authors = paper.get('authors', [])
                author_str = ', '.join(authors[:3]) + ('等' if len(authors) > 3 else '')
                year = paper.get('year', '未知年份')
                journal = paper.get('journal', '未知期刊')
                papers_info += f"{i+1}. {paper.get('title', '无标题')} - {author_str} ({year}) {journal}\n"

        prompt = f"""- Role: 学术综述选题专家
- Background: 用户已经收集了特定研究领域的文献资料，需要基于这些真实文献提供具体的综述选题建议。
- Profile: 你是一位在学术写作和文献综述方面有着丰富经验的专家，擅长从大量文献中识别研究热点、发现研究空白、提出有价值的综述选题。
- Skills: 你具备文献分析能力、选题策划能力、学术写作指导能力，能够基于真实文献数据提供具体可行的综述选题建议。
- Goals:
  1. 基于提供的文献分析该领域的研究热点和发展趋势
  2. 识别适合综述的具体角度和切入点
  3. 提供3-5个具体的综述选题建议
  4. 为每个选题提供详细的写作思路和文献支撑
- Constrains: 选题建议必须基于提供的真实文献，确保可行性和学术价值。避免过于宽泛或过于狭窄的选题。
- OutputFormat: 请用中文回答，使用HTML格式，包含适当的标题和段落标签。

现在请基于"{query}"领域的{len(selected_titles)}篇文献，提供具体的综述选题建议：

文献标题列表：
{titles_text}
{papers_info}

请提供：
1. 该领域研究热点分析
2. 3-5个具体的综述选题建议
3. 每个选题的写作思路和预期贡献
4. 推荐的文献组织方式

确保选题具有学术价值、可操作性和创新性。"""

        try:
            # 调用DeepSeek API
            suggestion_result = call_deepseek_api(prompt)

            return jsonify({
                'success': True,
                'suggestion': suggestion_result,
                'analyzed_papers': len(selected_titles),
                'total_papers': len(titles) if titles else len(papers)
            })

        except Exception as e:
            logger.error(f"DeepSeek API调用失败: {str(e)}")
            # 提供备用建议
            fallback_suggestion = f"""
            <h3>📝 {query} - 综述选题建议</h3>

            <h4>🔥 研究热点分析</h4>
            <p>基于{len(selected_titles)}篇文献分析，该领域当前的研究热点集中在技术创新、临床应用和方法学改进等方面。</p>

            <h4>📋 推荐综述选题</h4>
            <ol>
                <li><strong>{query}技术发展与临床应用综述</strong>
                    <p>梳理该领域的技术演进历程，总结当前主流技术的优缺点，分析临床应用现状和前景。</p>
                </li>
                <li><strong>{query}相关方法学比较与评价</strong>
                    <p>系统比较不同研究方法的适用性、准确性和局限性，为研究者选择合适方法提供参考。</p>
                </li>
                <li><strong>{query}领域的挑战与未来发展方向</strong>
                    <p>分析当前面临的主要技术和临床挑战，预测未来发展趋势和研究重点。</p>
                </li>
            </ol>

            <h4>💡 写作建议</h4>
            <p>建议采用系统性综述的方法，按照技术发展时间线或应用领域进行文献组织，确保综述的逻辑性和完整性。</p>

            <p><em>💡 注：由于AI分析服务暂时不可用，以上为基于文献数量和研究领域的基础建议。</em></p>
            """

            return jsonify({
                'success': True,
                'suggestion': fallback_suggestion,
                'analyzed_papers': len(selected_titles),
                'total_papers': len(titles) if titles else len(papers)
            })

    except Exception as e:
        logger.error(f"综述选题建议生成出错: {str(e)}\n{traceback.format_exc()}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/export/excel', methods=['POST'])
def export_excel():
    """导出Excel文件API端点"""
    try:
        data = request.get_json()
        query = data.get('query', '').strip()

        # 获取会话ID
        session_id = request.headers.get('sid') or session.get('session_id')
        if not session_id:
            return jsonify({'error': '无效的会话ID'}), 400

        # 从缓存中获取文献数据
        with papers_cache_lock:
            if session_id not in papers_cache:
                return jsonify({'error': '文献数据不存在，请重新检索'}), 404

            papers_data = papers_cache[session_id].copy()

        # 生成Excel文件
        file_content, filename = generate_excel_in_memory(papers_data)

        if file_content:
            # 创建BytesIO对象
            excel_buffer = BytesIO(file_content)
            excel_buffer.seek(0)
            return send_file(
                excel_buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
        else:
            return jsonify({'error': '生成Excel文件失败'}), 500

    except Exception as e:
        logger.error(f"导出Excel文件时发生错误: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/export/word', methods=['POST'])
def export_word():
    """导出Word文件API端点"""
    try:
        data = request.get_json()
        query = data.get('query', '').strip()

        # 获取会话ID
        session_id = request.headers.get('sid') or session.get('session_id')
        if not session_id:
            return jsonify({'error': '无效的会话ID'}), 400

        # 从缓存中获取文献数据
        with papers_cache_lock:
            if session_id not in papers_cache:
                return jsonify({'error': '文献数据不存在，请重新检索'}), 404

            papers_data = papers_cache[session_id].copy()

        # 生成Word文件
        file_content, filename = generate_word_in_memory(papers_data)

        if file_content:
            # 创建BytesIO对象
            word_buffer = BytesIO(file_content)
            word_buffer.seek(0)
            return send_file(
                word_buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            return jsonify({'error': '生成Word文件失败'}), 500

    except Exception as e:
        logger.error(f"导出Word文件时发生错误: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/research_topic_suggestion', methods=['POST'])
def research_topic_suggestion():
    """基于真实文献的论著选题建议API端点"""
    try:
        data = request.get_json()
        query = data.get('query', '').strip()
        titles = data.get('titles', [])
        papers = data.get('papers', [])

        if not query:
            return jsonify({
                'success': False,
                'error': '请提供研究方向'
            }), 400

        # 如果没有提供titles和papers，从缓存中获取
        if not titles and not papers:
            session_id = request.headers.get('sid') or session.get('session_id')
            if session_id:
                with papers_cache_lock:
                    if session_id in papers_cache:
                        papers_data = papers_cache[session_id]
                        papers = papers_data.get('papers', [])
                        titles = [paper.get('title', '') for paper in papers]
                        query = papers_data.get('query', query)  # 使用缓存中的查询词

        if not titles and not papers:
            return jsonify({
                'success': False,
                'error': '没有文献数据用于分析，请先进行文献检索'
            }), 400

        # 使用前150篇相关性最强的文献
        selected_titles = titles[:150] if titles else []
        selected_papers = papers[:150] if papers else []

        # 构建论著选题分析提示词
        titles_text = '\n'.join([f"{i+1}. {title}" for i, title in enumerate(selected_titles)])

        # 如果有完整的文献信息，提取更多细节
        papers_info = ""
        if selected_papers:
            papers_info = "\n\n文献详细信息（前20篇）：\n"
            for i, paper in enumerate(selected_papers[:20]):
                authors = paper.get('authors', [])
                author_str = ', '.join(authors[:3]) + ('等' if len(authors) > 3 else '')
                year = paper.get('year', '未知年份')
                journal = paper.get('journal', '未知期刊')
                papers_info += f"{i+1}. {paper.get('title', '无标题')} - {author_str} ({year}) {journal}\n"

        prompt = f"""- Role: 学术论著选题专家
- Background: 用户已经收集了特定研究领域的文献资料，需要基于这些真实文献发现研究空白，提供具体的论著选题建议。
- Profile: 你是一位在学术研究和论文写作方面有着丰富经验的专家，擅长从文献中识别研究空白、发现创新点、提出有价值的研究问题。
- Skills: 你具备文献分析能力、研究设计能力、创新思维能力，能够基于真实文献数据识别研究机会和提出可行的研究方案。
- Goals:
  1. 基于提供的文献识别该领域的研究空白和不足
  2. 发现具有创新性和可行性的研究问题
  3. 提供3-5个具体的论著选题建议
  4. 为每个选题提供研究思路和方法建议
- Constrains: 选题建议必须基于提供的真实文献，确保创新性、可行性和学术价值。避免重复已有研究或过于困难的选题。
- OutputFormat: 请用中文回答，使用HTML格式，包含适当的标题和段落标签。

现在请基于"{query}"领域的{len(selected_titles)}篇文献，提供具体的论著选题建议：

文献标题列表：
{titles_text}
{papers_info}

请提供：
1. 该领域研究空白分析
2. 3-5个具体的论著选题建议
3. 每个选题的研究思路和方法建议
4. 预期的创新点和学术贡献

确保选题具有创新性、可行性和实用价值。"""

        try:
            # 调用DeepSeek API
            suggestion_result = call_deepseek_api(prompt)

            return jsonify({
                'success': True,
                'suggestion': suggestion_result,
                'analyzed_papers': len(selected_titles),
                'total_papers': len(titles) if titles else len(papers)
            })

        except Exception as e:
            logger.error(f"DeepSeek API调用失败: {str(e)}")
            # 提供备用建议
            fallback_suggestion = f"""
            <h3>🔬 {query} - 论著选题建议</h3>

            <h4>🔍 研究空白分析</h4>
            <p>基于{len(selected_titles)}篇文献分析，该领域在方法创新、临床验证、机制探索等方面仍存在研究空白。</p>

            <h4>📋 推荐论著选题</h4>
            <ol>
                <li><strong>改进的{query}检测方法研究</strong>
                    <p>针对现有方法的局限性，开发更准确、更高效的检测技术。</p>
                </li>
                <li><strong>{query}在特定人群中的应用研究</strong>
                    <p>在特殊人群（如老年人、儿童等）中验证技术的有效性和安全性。</p>
                </li>
                <li><strong>{query}相关生物标志物研究</strong>
                    <p>发现和验证新的生物标志物，提高诊断准确性。</p>
                </li>
                <li><strong>{query}多中心临床研究</strong>
                    <p>开展大样本、多中心研究，提高结果的可靠性和推广性。</p>
                </li>
            </ol>

            <p><em>💡 注：由于AI分析服务暂时不可用，以上为基于文献数量和研究领域的专业基础建议。</em></p>
            """

            return jsonify({
                'success': True,
                'suggestion': fallback_suggestion,
                'analyzed_papers': len(selected_titles),
                'total_papers': len(titles) if titles else len(papers)
            })

    except Exception as e:
        logger.error(f"论著选题建议出错: {str(e)}\n{traceback.format_exc()}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/generate_full_review', methods=['POST'])
def generate_full_review():
    """生成完整综述文档API端点"""
    try:
        data = request.get_json()
        query = data.get('query', '').strip()
        titles = data.get('titles', [])
        papers = data.get('papers', [])

        if not query:
            return jsonify({
                'success': False,
                'error': '请提供研究方向'
            }), 400

        # 如果没有提供titles和papers，从缓存中获取
        if not titles and not papers:
            session_id = request.headers.get('sid') or session.get('session_id')
            if session_id:
                with papers_cache_lock:
                    if session_id in papers_cache:
                        papers_data = papers_cache[session_id]
                        papers = papers_data.get('papers', [])
                        titles = [paper.get('title', '') for paper in papers]
                        query = papers_data.get('query', query)  # 使用缓存中的查询词

        if not titles and not papers:
            return jsonify({
                'success': False,
                'error': '没有文献数据用于分析，请先进行文献检索'
            }), 400

        # 使用前150篇相关性最强的文献
        selected_titles = titles[:150] if titles else []
        selected_papers = papers[:150] if papers else []

        # 构建完整综述生成提示词
        titles_text = '\n'.join([f"{i+1}. {title}" for i, title in enumerate(selected_titles)])

        # 如果有完整的文献信息，提取更多细节
        papers_info = ""
        if selected_papers:
            papers_info = "\n\n文献详细信息（前50篇）：\n"
            for i, paper in enumerate(selected_papers[:50]):
                authors = paper.get('authors', [])
                author_str = ', '.join(authors[:3]) + ('等' if len(authors) > 3 else '')
                year = paper.get('year', '未知年份')
                journal = paper.get('journal', '未知期刊')
                abstract = paper.get('abstract', '')[:200] + '...' if paper.get('abstract') else '无摘要'
                papers_info += f"{i+1}. {paper.get('title', '无标题')} - {author_str} ({year}) {journal}\n   摘要: {abstract}\n\n"

        prompt = f"""- Role: 学术综述写作专家
- Background: 用户需要基于收集的真实文献生成一份完整的学术综述文档，要求内容全面、结构清晰、学术规范。
- Profile: 你是一位在学术写作和文献综述方面有着丰富经验的专家，擅长整合大量文献信息，撰写高质量的综述文章。
- Skills: 你具备文献综合分析能力、学术写作能力、逻辑思维能力，能够基于真实文献数据撰写结构完整、内容丰富的综述文档。
- Goals:
  1. 基于提供的文献撰写完整的综述文档
  2. 确保内容结构清晰、逻辑严密
  3. 包含所有综述必要的组成部分
  4. 提供专业的学术观点和见解
- Constrains: 综述内容必须基于提供的真实文献，确保学术严谨性和可信度。遵循学术写作规范。
- OutputFormat: 请用中文回答，使用HTML格式，包含适当的标题和段落标签，确保格式规范。

现在请基于"{query}"领域的{len(selected_titles)}篇文献，生成完整的综述文档：

文献标题列表：
{titles_text}
{papers_info}

请按以下结构撰写完整综述：

1. **研究背景与意义**
   - 领域发展历程
   - 研究重要性和临床意义
   - 当前面临的挑战

2. **文献检索策略**
   - 检索数据库和时间范围
   - 检索关键词和策略
   - 文献筛选标准

3. **研究现状分析**
   - 主要研究成果梳理
   - 技术方法发展现状
   - 临床应用进展

4. **技术方法比较**
   - 不同方法的优缺点
   - 适用场景分析
   - 性能指标比较

5. **存在问题与挑战**
   - 当前技术局限性
   - 临床应用障碍
   - 标准化问题

6. **发展趋势与展望**
   - 未来发展方向
   - 新兴技术趋势
   - 临床应用前景

7. **结论与建议**
   - 主要结论总结
   - 实践建议
   - 未来研究方向

确保内容专业、全面、具有学术价值。"""

        try:
            # 调用DeepSeek API
            review_result = call_deepseek_api(prompt)

            return jsonify({
                'success': True,
                'review': review_result,
                'analyzed_papers': len(selected_titles),
                'total_papers': len(titles) if titles else len(papers)
            })

        except Exception as e:
            logger.error(f"DeepSeek API调用失败: {str(e)}")
            # 提供备用综述
            fallback_review = f"""
            <h2>📚 {query} - 综述文档</h2>

            <h3>1. 研究背景与意义</h3>
            <p>{query}是当前医学研究的重要领域，具有重要的临床意义和应用价值。基于{len(selected_titles)}篇相关文献的分析，该领域在近年来取得了显著进展。</p>

            <h3>2. 文献检索策略</h3>
            <p>本综述检索了PubMed等主要医学数据库，时间范围为2020-2025年，共纳入{len(selected_titles)}篇高质量文献进行分析。</p>

            <h3>3. 研究现状分析</h3>
            <p>当前{query}领域的研究主要集中在技术创新、临床应用和方法学改进等方面，取得了一系列重要成果。</p>

            <h3>4. 技术方法比较</h3>
            <p>不同的技术方法各有优缺点，在准确性、效率和适用性方面存在差异，需要根据具体应用场景选择合适的方法。</p>

            <h3>5. 存在问题与挑战</h3>
            <p>该领域仍面临技术标准化、临床验证、成本控制等挑战，需要进一步的研究和改进。</p>

            <h3>6. 发展趋势与展望</h3>
            <p>未来{query}领域将朝着更加精准、智能、个性化的方向发展，有望在临床实践中发挥更大作用。</p>

            <h3>7. 结论与建议</h3>
            <p>建议加强多学科合作，推进技术创新，完善标准规范，促进{query}技术的临床转化和应用。</p>

            <p><em>💡 注：由于AI分析服务暂时不可用，以上为基于文献数量和研究领域的专业基础综述框架。</em></p>
            """

            return jsonify({
                'success': True,
                'review': fallback_review,
                'analyzed_papers': len(selected_titles),
                'total_papers': len(titles) if titles else len(papers)
            })

    except Exception as e:
        logger.error(f"完整综述生成出错: {str(e)}\n{traceback.format_exc()}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/test-keywords', methods=['POST'])
def test_keywords():
    """测试关键词提取功能"""
    try:
        data = request.json
        query = data.get('query', '')

        if not query:
            return jsonify({
                'success': False,
                'error': '请提供检索词'
            }), 400

        # 初始化分析器
        analyzer = JournalAnalyzer()

        # 获取文章数据
        articles = analyzer.fetch_journal_articles(query)

        if not articles:
            return jsonify({
                'success': False,
                'error': '未找到相关文章'
            }), 404

        # 测试关键词提取
        test_results = analyzer.test_extract_keywords(articles)

        # 生成词云数据（仅使用关键词）
        wordcloud_data = analyzer.extract_keywords(articles)[:50]  # 限制50个关键词

        # 准备返回数据
        response_data = {
            'success': True,
            'test_results': test_results,
            'wordcloud_data': wordcloud_data,
            'total_articles': len(articles)
        }
        
        logger.info(f"测试完成，数据预览:")
        logger.info(f"从关键词字段提取: {len(test_results['from_keywords'])} 个关键词")
        logger.info(f"从标题中提取: {len(test_results['from_titles'])} 个关键词")
        logger.info(f"重叠关键词: {len(test_results['overlap'])} 个")
        logger.info(f"词云数据: {len(wordcloud_data)} 个关键词")
        
        return jsonify(response_data)
        
    except Exception as e:
        logger.error(f"关键词提取测试失败: {str(e)}\n{traceback.format_exc()}")
        return jsonify({
            'success': False,
            'error': f'测试失败: {str(e)}'
        }), 500

# 解析推荐结果的辅助函数
def parse_recommendations_to_structured_data(recommendations_text, original_papers):
    """将推荐文本解析为结构化数据"""
    try:
        logger.info(f"开始解析推荐结果，原始文献数量: {len(original_papers)}")

        # 首先检查文本是否有换行符，如果没有，尝试用其他分隔符分割
        if '\n' not in recommendations_text:
            # 尝试用 --- 分割文献
            parts = recommendations_text.split('---')
            logger.info(f"文本没有换行符，用---分割得到 {len(parts)} 部分")
        else:
            parts = [recommendations_text]

        recommended_papers = []

        # 使用换行符分割文本进行解析
        import re
        lines = recommendations_text.split('\n')
        recommended_papers = []

        current_paper = None
        current_section = None
        content_buffer = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 检测文献标题行
            if line.startswith(('1. **', '2. **', '3. **', '4. **', '5. **')):
                # 保存上一篇文献
                if current_paper:
                    if current_section == 'reason' and content_buffer:
                        current_paper['reason'] = ' '.join(content_buffer).strip()
                    recommended_papers.append(current_paper)

                # 开始新文献
                title_match = re.search(r'\d+\.\s*\*\*(.*?)\*\*', line)
                if title_match:
                    title = title_match.group(1).strip()
                    current_paper = {
                        'title': title,
                        'chinese_title': '',
                        'authors': '',
                        'abstract': '',
                        'reason': ''
                    }
                    current_section = None
                    content_buffer = []
                    logger.info(f"找到文献: {title}")

            # 检测中文标题（方括号格式）
            elif line.startswith('[') and line.endswith(']') and current_paper:
                current_paper['chinese_title'] = line[1:-1]  # 去掉方括号
                logger.info(f"找到中文标题: {current_paper['chinese_title']}")

            # 检测各个部分
            elif line.startswith('**作者信息**'):
                current_section = 'authors'
                content_buffer = []
            elif line.startswith('**期刊信息**'):
                if current_section == 'authors' and content_buffer:
                    current_paper['authors'] = ' '.join(content_buffer).strip()
                current_section = 'journal'
                content_buffer = []
            elif line.startswith('**中文摘要总结**') or line.startswith('**摘要**'):
                current_section = 'abstract'
                content_buffer = []
            elif line.startswith('**推荐理由**'):
                if current_section == 'abstract' and content_buffer:
                    current_paper['abstract'] = ' '.join(content_buffer).strip()
                current_section = 'reason'
                content_buffer = []
            elif line.startswith('---'):
                if current_section == 'reason' and content_buffer:
                    current_paper['reason'] = ' '.join(content_buffer).strip()
                current_section = None
                content_buffer = []
            elif line.startswith('📊') or line.startswith('🎯'):
                if current_section == 'reason' and content_buffer:
                    current_paper['reason'] = ' '.join(content_buffer).strip()
                break
            elif current_section and not line.startswith('**'):
                content_buffer.append(line)

        # 处理最后一篇文献
        if current_paper:
            if current_section == 'reason' and content_buffer:
                current_paper['reason'] = ' '.join(content_buffer).strip()
            elif current_section == 'abstract' and content_buffer:
                current_paper['abstract'] = ' '.join(content_buffer).strip()
            recommended_papers.append(current_paper)

        logger.info(f"解析出 {len(recommended_papers)} 篇推荐文献")

        # 匹配原始文献数据
        structured_papers = []
        for i, rec_paper in enumerate(recommended_papers[:5]):
            title = rec_paper.get('title', '').strip()
            logger.info(f"正在匹配第 {i+1} 篇文献: {title}")

            if title:
                # 在原始文献中查找匹配的文献
                best_match = None
                best_score = 0

                for orig_paper in original_papers:
                    orig_title = orig_paper.get('title', '').strip()

                    # 计算标题相似度
                    title_lower = title.lower()
                    orig_title_lower = orig_title.lower()

                    # 完全匹配
                    if title_lower == orig_title_lower:
                        best_match = orig_paper
                        best_score = 100
                        break

                    # 包含匹配
                    if title_lower in orig_title_lower or orig_title_lower in title_lower:
                        score = max(len(title_lower), len(orig_title_lower)) / min(len(title_lower), len(orig_title_lower))
                        if score > best_score:
                            best_match = orig_paper
                            best_score = score

                    # 关键词匹配
                    title_words = set(title_lower.split())
                    orig_words = set(orig_title_lower.split())
                    common_words = title_words.intersection(orig_words)
                    if len(common_words) >= 3:  # 至少3个共同词
                        score = len(common_words) / max(len(title_words), len(orig_words))
                        if score > best_score:
                            best_match = orig_paper
                            best_score = score

                if best_match:
                    logger.info(f"找到匹配文献，相似度: {best_score}")
                    structured_paper = {
                        'title': best_match.get('title', ''),
                        'chinese_title': rec_paper.get('chinese_title', ''),
                        'authors': best_match.get('authors', []),
                        'abstract': rec_paper.get('abstract', best_match.get('abstract', '')),  # 优先使用AI生成的中文摘要
                        'journal': best_match.get('journal', ''),
                        'pub_year': best_match.get('pub_year', ''),
                        'url': best_match.get('url', ''),
                        'journal_info': best_match.get('journal_info', {}),
                        'impact_factor': best_match.get('impact_factor', ''),
                        'jcr_quartile': best_match.get('jcr_quartile', ''),
                        'cas_quartile': best_match.get('cas_quartile', ''),
                        'relevance': best_match.get('relevance', ''),
                        'recommendation_reason': rec_paper.get('reason', '')
                    }
                    structured_papers.append(structured_paper)
                else:
                    logger.warning(f"未找到匹配的原始文献: {title}")

        logger.info(f"成功匹配 {len(structured_papers)} 篇文献")
        return structured_papers

    except Exception as e:
        logger.error(f"解析推荐结果时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return []

def parse_precise_search_results(analysis_text, original_papers):
    """将精确查找结果解析为结构化数据"""
    try:
        logger.info(f"开始解析精确查找结果，原始文献数量: {len(original_papers)}")

        # 使用换行符分割文本进行解析
        import re
        lines = analysis_text.split('\n')
        found_papers = []

        current_paper = None
        current_section = None
        content_buffer = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 检测文献标题行（数字开头的标题）
            if re.match(r'^\d+\.\s*\*\*.*\*\*', line):
                # 保存上一篇文献
                if current_paper:
                    if current_section == 'selection_reason' and content_buffer:
                        current_paper['selection_reason'] = ' '.join(content_buffer).strip()
                    found_papers.append(current_paper)

                # 开始新文献
                title_match = re.search(r'\d+\.\s*\*\*(.*?)\*\*', line)
                if title_match:
                    title = title_match.group(1).strip()
                    current_paper = {
                        'title': title,
                        'chinese_title': '',
                        'authors': '',
                        'abstract': '',
                        'selection_reason': '',
                        'sufficiency_score': ''
                    }
                    current_section = None
                    content_buffer = []
                    logger.info(f"找到文献: {title}")

            # 检测中文标题（方括号格式）
            elif line.startswith('[') and line.endswith(']') and current_paper:
                current_paper['chinese_title'] = line[1:-1]  # 去掉方括号
                logger.info(f"找到中文标题: {current_paper['chinese_title']}")

            # 检测各个部分
            elif line.startswith('**作者：**') or line.startswith('**作者信息**'):
                current_section = 'authors'
                content_buffer = []
            elif line.startswith('**发表年份：**') or line.startswith('**期刊信息**'):
                if current_section == 'authors' and content_buffer:
                    current_paper['authors'] = ' '.join(content_buffer).strip()
                current_section = 'journal'
                content_buffer = []
            elif line.startswith('**符合条件充分性：**') or line.startswith('**充分性评分：**'):
                # 解析充分性评分
                import re
                score_match = re.search(r'(\d+)%', line)
                if score_match:
                    current_paper['sufficiency_score'] = score_match.group(1)
                current_section = None
                content_buffer = []
            elif line.startswith('**中文摘要总结**') or line.startswith('**摘要**'):
                current_section = 'abstract'
                content_buffer = []
            elif line.startswith('**符合筛选的理由**') or line.startswith('**筛选理由**'):
                if current_section == 'abstract' and content_buffer:
                    current_paper['abstract'] = ' '.join(content_buffer).strip()
                current_section = 'selection_reason'
                content_buffer = []
            elif line.startswith('---'):
                if current_section == 'selection_reason' and content_buffer:
                    current_paper['selection_reason'] = ' '.join(content_buffer).strip()
                current_section = None
                content_buffer = []
            elif current_section and not line.startswith('**'):
                content_buffer.append(line)

        # 处理最后一篇文献
        if current_paper:
            if current_section == 'selection_reason' and content_buffer:
                current_paper['selection_reason'] = ' '.join(content_buffer).strip()
            elif current_section == 'abstract' and content_buffer:
                current_paper['abstract'] = ' '.join(content_buffer).strip()
            found_papers.append(current_paper)

        logger.info(f"解析出 {len(found_papers)} 篇精确查找文献")

        # 匹配原始文献数据
        structured_papers = []
        for i, found_paper in enumerate(found_papers):
            title = found_paper.get('title', '').strip()
            logger.info(f"正在匹配第 {i+1} 篇文献: {title}")

            if title:
                # 在原始文献中查找匹配的文献
                best_match = None
                best_score = 0

                for orig_paper in original_papers:
                    orig_title = orig_paper.get('title', '').strip()

                    # 计算标题相似度
                    title_lower = title.lower()
                    orig_title_lower = orig_title.lower()

                    # 完全匹配
                    if title_lower == orig_title_lower:
                        best_match = orig_paper
                        best_score = 100
                        break

                    # 包含匹配
                    if title_lower in orig_title_lower or orig_title_lower in title_lower:
                        score = max(len(title_lower), len(orig_title_lower)) / min(len(title_lower), len(orig_title_lower))
                        if score > best_score:
                            best_match = orig_paper
                            best_score = score

                    # 关键词匹配
                    title_words = set(title_lower.split())
                    orig_words = set(orig_title_lower.split())
                    common_words = title_words.intersection(orig_words)
                    if len(common_words) >= 3:  # 至少3个共同词
                        score = len(common_words) / max(len(title_words), len(orig_words))
                        if score > best_score:
                            best_match = orig_paper
                            best_score = score

                if best_match:
                    logger.info(f"找到匹配文献，相似度: {best_score}")
                    structured_paper = {
                        'title': best_match.get('title', ''),
                        'chinese_title': found_paper.get('chinese_title', ''),
                        'authors': best_match.get('authors', []),
                        'abstract': found_paper.get('abstract', best_match.get('abstract', '')),  # 优先使用AI生成的中文摘要
                        'journal': best_match.get('journal', ''),
                        'pub_year': best_match.get('pub_year', ''),
                        'url': best_match.get('url', ''),
                        'journal_info': best_match.get('journal_info', {}),
                        'impact_factor': best_match.get('impact_factor', ''),
                        'jcr_quartile': best_match.get('jcr_quartile', ''),
                        'cas_quartile': best_match.get('cas_quartile', ''),
                        'relevance': best_match.get('relevance', ''),
                        'selection_reason': found_paper.get('selection_reason', ''),
                        'sufficiency_score': found_paper.get('sufficiency_score', '')
                    }
                    structured_papers.append(structured_paper)
                else:
                    logger.warning(f"未找到匹配的原始文献: {title}")

        logger.info(f"成功匹配 {len(structured_papers)} 篇文献")
        return structured_papers

    except Exception as e:
        logger.error(f"解析精确查找结果时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return []

# 智能问题推荐API端点

@app.route('/api/recommend_representative_papers', methods=['POST'])
def recommend_representative_papers():
    """推荐代表性文章"""
    try:
        # 从缓存中获取文献数据
        session_id = request.headers.get('sid') or session.get('session_id')
        papers = []

        if session_id:
            with papers_cache_lock:
                if session_id in papers_cache:
                    papers_data = papers_cache[session_id]
                    papers = papers_data.get('papers', [])

        if not papers:
            return jsonify({
                'success': False,
                'error': '没有文献数据用于分析，请先进行文献检索'
            }), 400

        # 选择相关度最高的80篇文献
        top_papers = sorted(papers, key=lambda x: x.get('relevance', 0), reverse=True)[:80]

        logger.info(f"开始推荐代表性文章，分析 {len(top_papers)} 篇高相关度文献")

        # 构建详细的文献信息
        literature_info = []
        for i, paper in enumerate(top_papers, 1):
            title = paper.get('title', '无标题')
            abstract = paper.get('abstract', '无摘要')
            journal_info = paper.get('journal_info', {})
            impact_factor = journal_info.get('impact_factor', 'N/A')
            jcr_quartile = journal_info.get('jcr_quartile', 'N/A')
            cas_quartile = journal_info.get('cas_quartile', 'N/A')
            relevance = paper.get('relevance', 0)

            # 格式化分区信息
            jcr_display = f"Q{jcr_quartile}" if jcr_quartile != 'N/A' else 'N/A'
            cas_display = f"{cas_quartile}区" if cas_quartile != 'N/A' else 'N/A'

            literature_info.append(f"""
文献{i}:
标题: {title}
摘要: {abstract}
影响因子: {impact_factor}
JCR分区: {jcr_display}
中科院分区: {cas_display}
相关度: {relevance:.1f}%
""")

        # 使用专业提示词
        expert_prompt = """
Role: 顶级科学家和文献评估专家
Background: 用户需要从大量文献中筛选出最具价值的5篇文章，这些文献包含了标题、摘要、分区以及影响因子等关键信息。用户希望从专业角度，依据影响因子、创新点和文章质量来筛选文献。
Profile: 你是一位在学术领域具有极高声望的顶级科学家，长期从事科研工作，对文献的评估有着丰富的经验和敏锐的洞察力。你能够快速准确地判断文献的创新性、学术价值以及其在领域内的影响力。
Skills: 你具备强大的文献分析能力，能够迅速从海量信息中提取关键内容；对不同学科领域的研究方法和评价标准有深入理解；能够综合考虑影响因子、创新点和文章质量等多个维度进行评估。
Goals: 从用户提供的文献中，依据影响因子、创新点和文章质量，筛选出最具代表性和价值的5篇文章。
Constrains: 仅从用户提供的文献信息中进行筛选，不进行额外的文献检索；确保推荐的文章在创新性、学术价值和影响力方面具有突出表现；推荐的文章需涵盖不同的研究方向和方法，以保证多样性和代表性。
OutputFormat: 输出格式应包括每篇推荐文献的标题、摘要、分区、影响因子以及推荐理由。
Workflow:
1. 对所有文献进行初步筛选，根据影响因子从高到低排序，优先考虑高分区和高影响因子的文献。
2. 深入分析每篇文献的摘要，评估其创新点和研究方法的独特性，筛选出具有显著创新性的文献。
3. 综合考虑文献的质量，包括研究设计、数据分析和结论的可靠性，最终确定5篇最具代表性的文献。
"""

        user_prompt = f"""
请从以下{len(top_papers)}篇文献中筛选出最具代表性和价值的5篇文章：

{chr(10).join(literature_info)}

**重要格式要求：**
1. 标题必须是中英文对照：英文原标题在第一行，中文翻译标题用方括号包围在第二行
2. 摘要必须是中文总结：用通俗易懂的中文总结文章主要内容，不要使用英文原文
3. 严格按照以下示例格式输出

**输出格式示例：**

1. **Long-Term Prognostic Utility of Coronary CT Angiography in Stable Patients With Diabetes Mellitus**
[长期预后价值：冠状动脉CT血管造影在稳定型糖尿病患者中的应用]

**期刊信息**
影响因子：12.800 | JCR分区：Q1 | 中科院分区：B1区 | 发表年份：2016

**中文摘要总结**
本研究评估了冠状动脉CT血管造影(CCTA)对糖尿病患者长期预后的预测价值。通过国际多中心CONFIRM注册研究，对1,823名糖尿病患者和1,823名非糖尿病患者进行5年随访。结果显示，无冠脉疾病的糖尿病患者死亡风险与非糖尿病患者相当；而非阻塞性和阻塞性冠脉疾病患者的死亡风险显著增加。特别值得注意的是，糖尿病患者即使仅有非阻塞性病变，其死亡风险也高于非糖尿病患者的阻塞性病变。研究证实CCTA对糖尿病患者具有重要的长期预后价值。

**推荐理由**
1) 创新点：首次大规模比较糖尿病患者与非糖尿病患者CCTA结果的长期预后差异
2) 学术价值：发表在心血管顶级期刊(JACC)，为糖尿病患者的冠脉评估提供重要循证依据
3) 影响力：研究结果改变了临床对糖尿病患者冠脉病变风险分层的认识
4) 方法学：采用国际多中心数据，严格的倾向匹配设计，随访时间长(5年)

---

2. **Prognostic Value of Coronary Computed Tomography Angiography in Diabetic Patients: A Meta-analysis**
[糖尿病患者冠状动脉CT血管造影预后价值的Meta分析]

**期刊信息**
影响因子：14.800 | JCR分区：Q1 | 中科院分区：B1区 | 发表年份：2016

**中文摘要总结**
这项Meta分析纳入8项研究共6,225名糖尿病患者，评估CCTA对心血管事件的预测价值。结果显示，阻塞性冠脉疾病(38%)和非阻塞性病变(36%)在糖尿病患者中普遍存在。阻塞性病变患者的年事件率高达17.1%，而非阻塞性病变和正常冠脉分别为4.5%和0.1%。CCTA可安全排除未来事件，并能识别需要强化治疗的高危患者。研究强调CCTA在糖尿病患者风险分层中的关键作用。

**推荐理由**
1) 创新点：首个针对糖尿病患者CCTA预后价值的系统评价
2) 学术价值：发表在糖尿病领域顶级期刊(Diabetes Care)，证据等级高
3) 影响力：为临床指南制定提供了重要依据
4) 方法学：采用严格的Meta分析方法，样本量大，结果可靠

---

**请严格按照上述格式输出5篇推荐文献：**

📚 代表性文章推荐

🏆 最具代表性的5篇文章

**摘要**
[完整摘要内容，保持原文格式]

**推荐理由**
[详细说明推荐理由]

---

4. **文献标题**
[完整文献标题]

**作者信息**
[作者列表]

**期刊信息**
影响因子：[数值] | JCR分区：[Q1/Q2/Q3/Q4] | 中科院分区：[1区/2区/3区/4区] | 发表年份：[年份]

**摘要**
[完整摘要内容，保持原文格式]

**推荐理由**
[详细说明推荐理由]

---

5. **文献标题**
[完整文献标题]

**作者信息**
[作者列表]

**期刊信息**
影响因子：[数值] | JCR分区：[Q1/Q2/Q3/Q4] | 中科院分区：[1区/2区/3区/4区] | 发表年份：[年份]

**摘要**
[完整摘要内容，保持原文格式]

**推荐理由**
[详细说明推荐理由]

---

📊 **筛选标准说明**
• 影响因子权重：30%（期刊影响力和声誉）
• 创新性权重：40%（研究方法和结果的创新程度）
• 学术质量权重：30%（研究设计、数据分析和结论可靠性）

🎯 **推荐总结**
**整体评价：** [对5篇推荐文献的整体评价]
**研究趋势：** [从推荐文献中观察到的研究趋势]
**建议关注：** [对研究者的具体建议]

*注：所有推荐文献均来自提供的{len(top_papers)}篇候选文献，按照专业评估标准进行筛选。*
"""

        # 调用DeepSeek API
        full_prompt = expert_prompt + "\n\n" + user_prompt
        recommendations = call_deepseek_api(full_prompt)

        # 在终端显示AI返回的完整内容，方便调试
        logger.info("=" * 80)
        logger.info("AI返回的推荐文献内容:")
        logger.info("=" * 80)
        logger.info(recommendations)
        logger.info("=" * 80)

        # 解析推荐结果为结构化数据
        structured_papers = parse_recommendations_to_structured_data(recommendations, top_papers)

        return jsonify({
            'success': True,
            'recommendations': recommendations,
            'structured_papers': structured_papers
        })

    except Exception as e:
        logger.error(f"推荐代表性文章时发生错误: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/precise_literature_search', methods=['POST'])
def precise_literature_search():
    """精确查找文献功能"""
    try:
        data = request.get_json()
        search_requirement = data.get('search_requirement', '')
        max_papers = data.get('max_papers', 100)

        if not search_requirement:
            return jsonify({
                'success': False,
                'error': '请提供查找要求'
            }), 400

        # 从缓存中获取文献数据
        session_id = request.headers.get('sid') or session.get('session_id')
        papers = []

        if session_id:
            with papers_cache_lock:
                if session_id in papers_cache:
                    papers_data = papers_cache[session_id]
                    papers = papers_data.get('papers', [])

        if not papers:
            return jsonify({
                'success': False,
                'error': '没有文献数据用于分析，请先进行文献检索'
            }), 400

        logger.info(f"开始精确查找文献，查找要求: {search_requirement}")
        logger.info(f"从 {len(papers)} 篇文献中筛选，最多返回 {max_papers} 篇")

        # 使用嵌入模型批量计算相关度
        try:
            logger.info(f"开始批量计算精确查找的嵌入相关度")

            # 使用现有的批量嵌入相关度计算函数
            batch_relevance_scores = calculate_batch_embedding_relevance(search_requirement, papers, EMBEDDING_BATCH_SIZE)

            if batch_relevance_scores:
                # 使用批量计算的结果
                papers_with_relevance = []
                for i, paper in enumerate(papers):
                    if i in batch_relevance_scores:
                        paper_copy = paper.copy()
                        paper_copy['relevance'] = batch_relevance_scores[i]
                        papers_with_relevance.append(paper_copy)

                # 按相关度排序并选择前N篇
                sorted_papers = sorted(papers_with_relevance, key=lambda x: x.get('relevance', 0), reverse=True)
                top_papers = sorted_papers[:max_papers]

                logger.info(f"批量计算完成，筛选出 {len(top_papers)} 篇高相关度文献")
            else:
                raise Exception("批量嵌入计算失败")

        except Exception as e:
            logger.error(f"计算文献相关度时出错: {e}")
            # 如果嵌入模型失败，使用关键词匹配作为备选方案
            top_papers = []
            search_keywords = search_requirement.lower().split()

            for paper in papers:
                title = paper.get('title', '').lower()
                abstract = paper.get('abstract', '').lower()
                paper_text = f"{title} {abstract}"

                # 计算关键词匹配度
                match_count = sum(1 for keyword in search_keywords if keyword in paper_text)
                relevance = (match_count / len(search_keywords)) * 100 if search_keywords else 0

                if relevance > 0:
                    paper_copy = paper.copy()
                    paper_copy['relevance'] = round(relevance, 1)
                    top_papers.append(paper_copy)

            # 按相关度排序并限制数量
            top_papers = sorted(top_papers, key=lambda x: x.get('relevance', 0), reverse=True)[:max_papers]

        if not top_papers:
            return jsonify({
                'success': True,
                'analysis': f'未找到与"{search_requirement}"相关的文献，请尝试调整查找要求。',
                'structured_papers': []
            })

        # 构建详细的文献信息用于AI分析
        literature_info = []
        for i, paper in enumerate(top_papers, 1):
            title = paper.get('title', '无标题')
            abstract = paper.get('abstract', '无摘要')
            journal_info = paper.get('journal_info', {})
            impact_factor = journal_info.get('impact_factor', 'N/A')
            jcr_quartile = journal_info.get('jcr_quartile', 'N/A')
            cas_quartile = journal_info.get('cas_quartile', 'N/A')
            relevance = paper.get('relevance', 0)

            # 格式化分区信息
            jcr_display = f"Q{jcr_quartile}" if jcr_quartile != 'N/A' else 'N/A'
            cas_display = f"{cas_quartile}区" if cas_quartile != 'N/A' else 'N/A'

            literature_info.append(f"""
文献{i}:
标题: {title}
摘要: {abstract}
影响因子: {impact_factor}
JCR分区: {jcr_display}
中科院分区: {cas_display}
相关度: {relevance:.1f}%
""")

        # 使用专业提示词进行严格筛选
        expert_prompt = f"""
Role: 严格的文献筛选专家和学术研究顾问

Background: 用户需要从大量文献中严格筛选出完全符合特定要求"{search_requirement}"的文献。这是一个精确筛选任务，要求极高的准确性和严格性。

Profile: 你是一位在学术研究领域具有极高专业水准的文献筛选专家，具备敏锐的判断力和严格的筛选标准。你能够准确识别文献的核心研究内容，严格区分相关和不相关的研究，绝不允许任何不符合要求的文献通过筛选。

Skills: 你具备：
1. 极强的文献内容分析能力，能够准确理解文献的核心研究对象和方法
2. 严格的逻辑判断能力，能够准确区分直接相关、间接相关和不相关的研究
3. 深厚的学科知识背景，能够准确识别不同研究领域的边界和区别

Goals: 从用户提供的文献中，严格筛选出完全符合"{search_requirement}"要求的文献，绝对排除任何不符合要求的文献。

**严格筛选标准：**
1. **直接相关性要求**：文献的核心研究内容必须直接涉及"{search_requirement}"
2. **排除标准**：严格排除以下情况的文献：
   - 仅在背景介绍中提及相关内容，但核心研究不是"{search_requirement}"
   - 研究对象虽然相关但属于不同类别（例如：要求降糖药物研究，则排除他汀类、降压药等其他药物研究）
   - 仅在讨论部分简单提及，但主要研究内容不符合要求
   - 综述文章中虽然涉及相关内容，但不是专门针对"{search_requirement}"的综述

3. **符合条件的充分性评估**：对于符合要求的文献，按以下标准评估其符合条件的充分性：
   - 高度符合（90-100%）：文献完全专注于"{search_requirement}"，是该领域的核心研究
   - 中度符合（70-89%）：文献主要研究"{search_requirement}"，但可能涉及部分相关但不同的内容
   - 基本符合（50-69%）：文献涉及"{search_requirement}"，但同时研究其他相关领域

Constrains:
1. 必须严格按照筛选标准执行，宁可漏选也不能错选
2. 只有真正符合"{search_requirement}"的文献才能被选中
3. 必须详细说明每篇文献符合筛选条件的具体理由和充分性程度

OutputFormat: 按符合条件的充分性从高到低排序，列出文献的中英文标题、作者、发表年份、期刊名称、中文摘要总结、符合条件的充分性评分以及详细的筛选理由。

**重要格式要求：**
1. 标题必须是中英文对照：英文原标题在第一行，中文翻译标题用方括号包围在第二行
2. 摘要必须是中文总结：用通俗易懂的中文总结文章主要内容
3. 必须包含符合条件的充分性评分（百分比）
4. 详细说明符合筛选条件的具体理由

**输出格式示例：**

1. **Efficacy and Safety of Novel Glucose-Lowering Agents in Type 2 Diabetes**
[新型降糖药物在2型糖尿病中的疗效和安全性研究]

**作者：** 张三, 李四, 王五
**发表年份：** 2023
**期刊名称：** 糖尿病研究
**影响因子：** 8.5 | **JCR分区：** Q1 | **中科院分区：** 1区
**符合条件充分性：** 95%

**中文摘要总结**
本研究专门评估了三种新型降糖药物（GLP-1受体激动剂、SGLT-2抑制剂、DPP-4抑制剂）在2型糖尿病患者中的降糖效果和安全性。通过随机对照试验，比较了这些药物的HbA1c降低幅度、低血糖风险和心血管安全性。

**符合筛选的理由**
该文献完全专注于降糖药物研究，核心内容就是评估新型降糖药物的疗效和安全性，与筛选要求"{search_requirement}"高度匹配。研究对象明确为降糖药物，研究目标明确为降糖效果评估，符合条件充分性达到95%。

---
"""

        user_prompt = f"""
请从以下{len(top_papers)}篇文献中严格筛选出完全符合"{search_requirement}"要求的文献：

{chr(10).join(literature_info)}

**严格筛选要求：**
1. **核心研究内容必须直接涉及"{search_requirement}"** - 不能仅在背景或讨论中提及
2. **严格排除不符合要求的文献** - 例如：
   - 如果要求是"降糖药物研究"，则必须排除他汀类药物、降压药物、抗凝药物等其他类型药物的研究
   - 如果要求是"人工智能医学影像"，则必须排除传统影像技术、AI在其他医学领域的应用
3. **宁可少选也不能错选** - 只有真正符合要求的文献才能被选中
4. **按符合条件的充分性排序** - 最符合要求的文献排在最前面，并且所有符合要求的文献都要输出

**输出要求：**
1. 每篇文献都有中英文对照标题
2. 摘要用中文总结且通俗易懂
3. 必须包含符合条件的充分性评分（百分比）
4. 详细说明为什么这篇文献完全符合"{search_requirement}"的要求
5. 按符合条件的充分性从高到低排序

**重要提醒：如果某篇文献不能明确判断是否完全符合"{search_requirement}"的要求，请直接排除，不要包含在结果中。**
"""

        # 调用DeepSeek API
        full_prompt = expert_prompt + "\n\n" + user_prompt
        analysis_result = call_deepseek_api(full_prompt)

        # 在终端显示AI返回的完整内容，方便调试
        logger.info("=" * 80)
        logger.info("AI返回的精确查找结果:")
        logger.info("=" * 80)
        logger.info(analysis_result)
        logger.info("=" * 80)

        # 解析结果为结构化数据
        structured_papers = parse_precise_search_results(analysis_result, top_papers)

        return jsonify({
            'success': True,
            'analysis': analysis_result,
            'structured_papers': structured_papers
        })

    except Exception as e:
        logger.error(f"精确查找文献时发生错误: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/suggest_further_search', methods=['POST'])
def suggest_further_search():
    """建议进一步检索"""
    try:
        data = request.get_json()
        query = data.get('query', '')

        if not query:
            return jsonify({
                'success': False,
                'error': '请提供研究方向'
            }), 400

        # 从缓存中获取文献数据
        session_id = request.headers.get('sid') or session.get('session_id')
        papers = []

        if session_id:
            with papers_cache_lock:
                if session_id in papers_cache:
                    papers_data = papers_cache[session_id]
                    papers = papers_data.get('papers', [])

        if not papers:
            return jsonify({
                'success': False,
                'error': '没有文献数据用于分析，请先进行文献检索'
            }), 400

        # 分析当前检索结果
        titles = [paper.get('title', '') for paper in papers[:50]]  # 使用前50篇

        prompt = f"""
基于当前检索到的{len(papers)}篇文献（查询："{query}"），请分析并提供进一步精准检索的建议：

当前检索结果的代表性文献标题：
{chr(10).join([f"{i+1}. {title}" for i, title in enumerate(titles[:20])])}

请按以下格式提供建议：

## 🔍 进一步检索建议

### 📊 当前检索结果分析：
- 文献数量：{len(papers)}篇
- 主要研究方向：[分析主要研究方向]
- 覆盖范围：[评估覆盖的研究范围]

### 🎯 精准检索建议：

#### 1. 细化检索策略
- **更具体的关键词组合**：[建议具体的关键词]
- **时间范围调整**：[建议合适的时间范围]
- **研究类型筛选**：[建议关注的研究类型]

#### 2. 扩展检索方向
- **相关子领域**：[建议探索的相关领域]
- **交叉学科**：[建议的交叉研究方向]
- **新兴技术应用**：[相关的新技术或方法]

#### 3. 具体检索建议
- **PubMed检索式**：[提供具体的检索式]
- **筛选条件**：[建议的筛选条件]
- **期刊范围**：[建议关注的期刊类型]

### 💡 检索优化提示：
- 如何提高检索精准度
- 如何避免遗漏重要文献
- 如何平衡查全率和查准率
"""

        # 调用AI分析
        suggestions = call_deepseek_api(prompt)

        return jsonify({
            'success': True,
            'suggestions': suggestions
        })

    except Exception as e:
        logger.error(f"建议进一步检索时发生错误: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/analyze_research_frontiers', methods=['POST'])
def analyze_research_frontiers():
    """分析前沿研究方向"""
    try:
        data = request.get_json()
        query = data.get('query', '')

        if not query:
            return jsonify({
                'success': False,
                'error': '请提供研究方向'
            }), 400

        # 从缓存中获取文献数据
        session_id = request.headers.get('sid') or session.get('session_id')
        papers = []

        if session_id:
            with papers_cache_lock:
                if session_id in papers_cache:
                    papers_data = papers_cache[session_id]
                    papers = papers_data.get('papers', [])

        if not papers:
            return jsonify({
                'success': False,
                'error': '没有文献数据用于分析，请先进行文献检索'
            }), 400

        # 分析最新的文献（按年份排序）
        recent_papers = sorted(papers, key=lambda x: x.get('pub_year', 0), reverse=True)[:30]
        titles = [paper.get('title', '') for paper in recent_papers]

        prompt = f"""
基于以下{len(titles)}篇最新文献，请分析该领域的前沿研究方向和发展趋势：

最新文献标题：
{chr(10).join([f"{i+1}. {title}" for i, title in enumerate(titles)])}

请按以下格式分析：

## 🚀 前沿研究方向分析

### 📈 当前研究热点：
- **热点方向1**：[描述具体的研究热点]
- **热点方向2**：[描述具体的研究热点]
- **热点方向3**：[描述具体的研究热点]

### 🔬 新兴技术趋势：
- **技术趋势1**：[新兴技术及其应用]
- **技术趋势2**：[新兴技术及其应用]
- **技术趋势3**：[新兴技术及其应用]

### 🎯 未来发展方向：
- **方向1**：[未来可能的发展方向]
- **方向2**：[未来可能的发展方向]
- **方向3**：[未来可能的发展方向]

### 💡 研究机会识别：
- **机会1**：[具体的研究机会和价值]
- **机会2**：[具体的研究机会和价值]
- **机会3**：[具体的研究机会和价值]

### 📊 趋势总结：
- 该领域的整体发展趋势
- 值得关注的新兴交叉领域
- 对研究者的建议
"""

        # 调用AI分析
        analysis = call_deepseek_api(prompt)

        return jsonify({
            'success': True,
            'analysis': analysis
        })

    except Exception as e:
        logger.error(f"分析前沿研究方向时发生错误: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/identify_research_gaps', methods=['POST'])
def identify_research_gaps():
    """识别研究空白与机会"""
    try:
        data = request.get_json()
        query = data.get('query', '')

        if not query:
            return jsonify({
                'success': False,
                'error': '请提供研究方向'
            }), 400

        # 从缓存中获取文献数据
        session_id = request.headers.get('sid') or session.get('session_id')
        papers = []

        if session_id:
            with papers_cache_lock:
                if session_id in papers_cache:
                    papers_data = papers_cache[session_id]
                    papers = papers_data.get('papers', [])

        if not papers:
            return jsonify({
                'success': False,
                'error': '没有文献数据用于分析，请先进行文献检索'
            }), 400

        # 分析高相关度文献
        top_papers = sorted(papers, key=lambda x: x.get('relevance', 0), reverse=True)[:40]
        titles = [paper.get('title', '') for paper in top_papers]

        prompt = f"""
基于以下{len(titles)}篇高相关度文献，请识别该领域的研究空白和潜在机会：

文献标题：
{chr(10).join([f"{i+1}. {title}" for i, title in enumerate(titles)])}

请按以下格式分析：

## 🔬 研究空白与机会识别

### 🕳️ 主要研究空白：

#### 1. 方法学空白
- **空白描述**：[具体的方法学缺陷或不足]
- **影响程度**：[对领域发展的影响]
- **填补建议**：[如何填补这个空白]

#### 2. 应用领域空白
- **空白描述**：[未充分探索的应用领域]
- **市场潜力**：[应用前景和价值]
- **研究建议**：[具体的研究方向]

#### 3. 理论基础空白
- **空白描述**：[理论层面的不足]
- **重要性**：[理论完善的必要性]
- **发展路径**：[理论发展的可能路径]

### 💎 潜在研究机会：

#### 1. 交叉学科机会
- **机会描述**：[跨学科研究的可能性]
- **创新潜力**：[可能产生的创新成果]
- **实施建议**：[如何开展交叉研究]

#### 2. 技术创新机会
- **机会描述**：[技术创新的空间]
- **技术路线**：[可能的技术发展路径]
- **预期成果**：[技术创新的预期效果]

#### 3. 临床转化机会
- **机会描述**：[从基础到临床的转化空间]
- **转化价值**：[临床应用的潜在价值]
- **转化策略**：[具体的转化路径]

### 🎯 优先研究建议：
- **高优先级项目**：[最值得优先开展的研究]
- **中期发展项目**：[中期可以考虑的研究方向]
- **长期探索项目**：[长期战略性研究方向]

### 💡 创新点提示：
- 如何在现有基础上实现突破
- 哪些技术组合可能产生创新
- 如何避免重复研究
"""

        # 调用AI分析
        analysis = call_deepseek_api(prompt)

        return jsonify({
            'success': True,
            'analysis': analysis
        })

    except Exception as e:
        logger.error(f"识别研究空白时发生错误: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

if __name__ == '__main__':
    try:
        # 检查NLTK数据
        nltk_data_path = os.path.join(os.path.expanduser('~'), 'nltk_data')
        
        if os.path.exists(nltk_data_path):
            logger.info("NLTK数据目录已存在，跳过下载")
        else:
            logger.info("首次运行，开始下载NLTK数据...")
            try:
                nltk.download('punkt', quiet=True)
                nltk.download('wordnet', quiet=True)
                logger.info("NLTK数据下载完成")
            except Exception as e:
                logger.error(f"NLTK数据下载失败: {str(e)}")
                logger.warning("将使用基础分词功能")
        
        # 检查必要的配置
        if not DEEPSEEK_API_KEY:
            raise ValueError("未设置DEEPSEEK_API_KEY")
        if not PUBMED_API_KEY:
            raise ValueError("未设置PUBMED_API_KEY")
        
        # 启动性能监控线程
        monitor_thread = threading.Thread(target=monitor_system_performance, daemon=True)
        monitor_thread.start()
        
        logger.info("正在启动应用服务器...")
        # 使用socketio启动应用
        socketio.run(app, debug=True, host='0.0.0.0', port=5000)
    except Exception as e:
        logger.error(f"启动失败: {str(e)}\n{traceback.format_exc()}")
        raise 